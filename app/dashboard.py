from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
import sqlite3
import ssl
import subprocess
import uuid
import ast
import math
from collections import Counter
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import pandas as pd
import httpx
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.encoders import jsonable_encoder
from fastapi.responses import FileResponse, Response
from openai import AzureOpenAI
from pydantic import BaseModel, EmailStr, Field, field_validator
from sqlalchemy import (
    Boolean,
    Date,
    DateTime,
    Float,
    ForeignKey,
    Integer,
    LargeBinary,
    String,
    Text,
    UniqueConstraint,
    create_engine,
    delete,
    func,
    inspect,
    select,
)
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, relationship, sessionmaker

IEEE_BLUE = "#00629B"
IEEE_TEAL = "#00843D"
IEEE_AMBER = "#FFB81C"
IEEE_ORANGE = "#E87722"
IEEE_RED = "#BA0C2F"
IEEE_GRAY = "#5B6770"

LIFECYCLE_PHASES = [
    "Unknown",
    "Expression of Interest",
    "Proposal Under Review",
    "ITSS Approved",
    "IEEE Application and MOU",
    "Detailed Planning",
    "Submission and Review",
    "Registration and Final Preparation",
    "Conference Delivery",
    "Proceedings Processing",
    "Financial and Administrative Closure",
    "Closed",
]
NORMALIZED_STATUSES = [
    "Unknown",
    "Open",
    "Not Started",
    "In Progress",
    "Submitted",
    "Awaiting IEEE",
    "Awaiting Conference",
    "Awaiting External Party",
    "Approved",
    "Complete",
    "Published",
    "Closed",
    "Blocked",
    "Rejected",
    "Cancelled",
    "Not Applicable",
]
CONFERENCE_SERIES = [
    ("ITSC", "International Conference on Intelligent Transportation Systems", True),
    ("IV", "Intelligent Vehicles Symposium", True),
    ("FISTS", "Forum for Innovative and Sustainable Transportation Systems", False),
    ("ICIRT", "International Conference on Intelligent Rail Transportation", False),
    ("ICVES", "International Conference on Vehicular Electronics and Safety", False),
    ("ISI", "International Conference on Intelligence and Security Informatics", False),
    ("SOLI", "International Conference on Service Operations and Logistics, and Informatics", False),
    ("MESA", "International Conference on Mechatronic and Embedded Systems and Applications", False),
    ("VNC", "Vehicular Networking Conference", False),
    ("Other TCS Conference", "Other TCS Conference", False),
    ("Custom Conference Series", "Custom Conference Series", False),
]
SPONSORSHIP_TYPES = [
    "Flagship",
    "Financially Sponsored",
    "Financially Co-Sponsored",
    "Technically Co-Sponsored",
]
CONTACT_ROLES = [
    "General Chair",
    "Program Chair",
    "Finance Chair",
    "Publications Chair",
    "Registration Chair",
    "Local Organization Chair",
    "Information Contact",
    "Conference Submitter",
    "IEEE Financial Analyst",
    "Other",
]
ISSUE_CATEGORIES = [
    "Governance",
    "Finance",
    "Publication",
    "Operations",
    "Registration",
    "Data Quality",
    "Document",
    "AI Review",
]
ISSUE_SEVERITIES = ["Informational", "Low", "Medium", "High", "Critical"]
REVIEW_ASSESSMENTS = [
    "Unreviewed",
    "On Track",
    "Needs Follow-up",
    "Not an Issue",
]
DERIVED_CONFERENCE_STATUSES = [
    "Not Started",
    "In Progress",
    "On Track",
    "Attention Needed",
    "At Risk",
    "Critical",
    "Blocked",
    "Complete",
    "Closed",
    "Cancelled",
]
SCORE_WEIGHTS = {
    "Governance and Approvals": 20.0,
    "Finance and Banking": 20.0,
    "Technical Program and Publication": 25.0,
    "Operations and Contracts": 15.0,
    "Registration and Event Readiness": 10.0,
    "Post-Conference Publication and Closure": 10.0,
}
REFERENCE_CONFIG_DEFAULTS = {
    "committee_members": ["Ahmed Hussein", "Daniel Medina"],
    "conference_series": [{"code": code, "name": name, "flagship": flagship} for code, name, flagship in CONFERENCE_SERIES],
    "lifecycle_phases": LIFECYCLE_PHASES,
    "conference_statuses": DERIVED_CONFERENCE_STATUSES,
    "normalized_statuses": NORMALIZED_STATUSES,
    "sponsorship_types": SPONSORSHIP_TYPES,
    "contact_roles": CONTACT_ROLES,
    "issue_categories": ISSUE_CATEGORIES,
    "issue_severities": ISSUE_SEVERITIES,
    "review_assessments": REVIEW_ASSESSMENTS,
}
REFERENCE_CONFIG_LABELS = {
    "committee_members": "ITSS Committee Members",
    "conference_series": "Conference Series",
    "lifecycle_phases": "Lifecycle Phases",
    "conference_statuses": "Conference Status Values",
    "normalized_statuses": "Status Values",
    "sponsorship_types": "Sponsorship Types",
    "contact_roles": "Contact Roles",
    "issue_categories": "Issue Categories",
    "issue_severities": "Issue Severities",
    "review_assessments": "Review Assessments",
}
MILESTONE_SEEDS = [
    ("APPLICATION", "Conference application approved", "Governance and Approvals", "Approved", -540),
    ("MOU", "MOU signed", "Governance and Approvals", "Approved", -450),
    ("BUDGET", "Budget approved", "Finance and Banking", "Approved", -365),
    ("BANKING", "Banking details", "Finance and Banking", "Complete", -240),
    ("CFP", "Call for papers published", "Technical Program and Publication", "Complete", -300),
    ("REVIEWS", "Paper review complete", "Technical Program and Publication", "Complete", -90),
    ("VENUE", "Venue and local arrangements confirmed", "Operations and Contracts", "Complete", -270),
    ("REGISTRATION", "Registration open", "Registration and Event Readiness", "Complete", -180),
    ("PROCEEDINGS", "Proceedings submitted", "Post-Conference Publication and Closure", "Submitted", 30),
    ("FIN_CLOSE", "Financial close complete", "Post-Conference Publication and Closure", "Closed", 180),
]

# Milestone date offsets as months and days relative to conference start date.
# Positive = after start date, negative = before start date.
MILESTONE_DATE_DEFAULTS = {
    # IEEE/ITSS conference lifecycle: application and MOU ~18-24 months before
    "APPLICATION": {"anchor": "start", "months": -24, "days": 0},
    "MOU": {"anchor": "start", "months": -18, "days": 0},
    # Budget and banking ~12 months before
    "BUDGET": {"anchor": "start", "months": -12, "days": 0},
    "BANKING": {"anchor": "start", "months": -9, "days": 0},
    # CFP ~10 months before, reviews ~3-4 months before
    "CFP": {"anchor": "start", "months": -10, "days": 0},
    "REVIEWS": {"anchor": "start", "months": -4, "days": 0},
    # Venue ~9-12 months before, registration ~4-6 months before
    "VENUE": {"anchor": "start", "months": -12, "days": 0},
    "REGISTRATION": {"anchor": "start", "months": -5, "days": 0},
    # Proceedings ~2-3 months after, financial close ~6-9 months after
    "PROCEEDINGS": {"anchor": "end", "months": 2, "days": 0},
    "FIN_CLOSE": {"anchor": "end", "months": 9, "days": 0},
}
DEFAULT_MILESTONE_STATUS_SCORES = {
    "completed": 100.0,
    "unknown": 0.0,
    "no_due_date": 65.0,
    "not_started_far": 80.0,
    "not_started_upcoming": 60.0,
    "not_started_due_soon": 30.0,
    "not_started_overdue": 0.0,
    "in_progress_on_time": 70.0,
    "in_progress_recently_overdue": 45.0,
    "in_progress_overdue": 20.0,
    "awaiting_on_time": 80.0,
    "awaiting_recently_overdue": 60.0,
    "awaiting_overdue": 35.0,
    "blocked": 0.0,
}

DEFAULT_SCORE_SETTINGS = {
    "dimension_weights": SCORE_WEIGHTS,
    "milestone_status_scores": DEFAULT_MILESTONE_STATUS_SCORES,
    "issue_severity_penalties": {"Critical": 12.0, "High": 6.0, "Medium": 3.0, "Low": 1.0, "Informational": 0.0},
    "issue_assessment_factors": {"Unreviewed": 1.0, "Needs Follow-up": 1.0, "On Track": 0.25, "Not an Issue": 0.0},
    "issue_penalty_cap": 40.0,
    "lateness_step_days": 30.0,
    "lateness_cap_factor": 3.0,
    "score_formula": "max(0, min(100, base_score - issue_penalty))",
}

DEFAULT_ASSISTANT_SYSTEM_PROMPT = (
    "You are the IEEE ITSS Conference Operations Assistant. Use the retrieved RAG knowledge-base excerpts, "
    "conference facts, milestones, comments, issues, contacts, and finance/publication/application data provided by the portal. "
    "Answer only from that evidence and clearly separate IEEE/ITSS policy facts from recommendations. Cite the source document "
    "title for every policy or guideline claim. If the knowledge base does not contain enough evidence, say what is missing and "
    "suggest the exact document or conference record that should be uploaded or updated. Do not invent dates, approvals, roles, "
    "or policy requirements."
)


def load_local_env(root: Path | None = None) -> None:
    root = root or Path(__file__).resolve().parents[1]
    env_path = root / ".env"
    if not env_path.exists():
        return
    for raw_line in env_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and (key not in os.environ or os.environ.get(key, "").strip() == ""):
            os.environ[key] = value


APP_ROOT = Path(__file__).resolve().parents[1]


def app_path(env_name: str, default: str) -> Path:
    path = Path(os.environ.get(env_name, default)).expanduser()
    return path if path.is_absolute() else APP_ROOT / path


class Base(DeclarativeBase):
    pass


class Conference(Base):
    __tablename__ = "conferences"
    __table_args__ = (
        UniqueConstraint("conference_number", name="uq_conference_number"),
        UniqueConstraint("normalized_acronym", "year", "parent_conference_id", name="uq_acronym_year_parent"),
    )

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_number: Mapped[str | None] = mapped_column(String(64), nullable=True)
    acronym: Mapped[str] = mapped_column(String(24), index=True)
    normalized_acronym: Mapped[str] = mapped_column(String(24), index=True)
    year: Mapped[int] = mapped_column(Integer, index=True)
    official_title: Mapped[str] = mapped_column(String(500))
    canonical_name: Mapped[str] = mapped_column(String(64), index=True)
    conference_series: Mapped[str] = mapped_column(String(160), index=True)
    conference_category: Mapped[str] = mapped_column(String(80), default="Portfolio")
    sponsorship_type: Mapped[str] = mapped_column(String(80), index=True)
    parent_conference_id: Mapped[str | None] = mapped_column(ForeignKey("conferences.id"), nullable=True)
    lifecycle_phase: Mapped[str] = mapped_column(String(120), default=LIFECYCLE_PHASES[0])
    suggested_phase: Mapped[str] = mapped_column(String(120), default=LIFECYCLE_PHASES[0])
    phase_override: Mapped[bool] = mapped_column(Boolean, default=False)
    conference_status: Mapped[str] = mapped_column(String(80), default="Unknown")
    active: Mapped[bool] = mapped_column(Boolean, default=True, index=True)
    start_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    end_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    submission_deadline: Mapped[date | None] = mapped_column(Date, nullable=True)
    notification_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    camera_ready_deadline: Mapped[date | None] = mapped_column(Date, nullable=True)
    city: Mapped[str | None] = mapped_column(String(120), nullable=True)
    state_province: Mapped[str | None] = mapped_column(String(120), nullable=True)
    country: Mapped[str | None] = mapped_column(String(120), nullable=True)
    ieee_region: Mapped[str | None] = mapped_column(String(40), nullable=True)
    venue: Mapped[str | None] = mapped_column(String(240), nullable=True)
    website: Mapped[str | None] = mapped_column(String(400), nullable=True)
    estimated_attendees: Mapped[int | None] = mapped_column(Integer, nullable=True)
    actual_attendees: Mapped[int | None] = mapped_column(Integer, nullable=True)
    last_reviewed_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    comments: Mapped[str | None] = mapped_column(Text, nullable=True)
    project_code: Mapped[str | None] = mapped_column(String(80), nullable=True)
    project_indicator: Mapped[str | None] = mapped_column(String(80), nullable=True)
    financial_analyst: Mapped[str | None] = mapped_column(String(160), nullable=True)
    committee_contact: Mapped[str | None] = mapped_column(String(160), nullable=True)
    application_status_raw: Mapped[str | None] = mapped_column(String(160), nullable=True)
    application_status: Mapped[str] = mapped_column(String(80), default="Unknown")
    application_submitted_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    application_approved_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    mou_status_raw: Mapped[str | None] = mapped_column(String(160), nullable=True)
    mou_status: Mapped[str] = mapped_column(String(80), default="Unknown")
    mou_signed_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    finance_status: Mapped[str] = mapped_column(String(80), default="Unknown")
    currency: Mapped[str | None] = mapped_column(String(12), nullable=True)
    total_income_current: Mapped[float | None] = mapped_column(Float, nullable=True)
    total_expense_current: Mapped[float | None] = mapped_column(Float, nullable=True)
    budgeted_income_total: Mapped[float | None] = mapped_column(Float, nullable=True)
    budgeted_expense_total: Mapped[float | None] = mapped_column(Float, nullable=True)
    itss_loan_requested: Mapped[bool] = mapped_column(Boolean, default=False)
    itss_loan_amount: Mapped[float | None] = mapped_column(Float, nullable=True)
    accounting_close_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    publication_status: Mapped[str] = mapped_column(String(80), default="Unknown")
    proceedings_submitted_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    xplore_posting_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    last_source_update: Mapped[date | None] = mapped_column(Date, nullable=True)
    source_details_json: Mapped[str] = mapped_column(Text, default="{}")
    score: Mapped[float] = mapped_column(Float, default=0.0)
    base_score: Mapped[float] = mapped_column(Float, default=0.0)
    issue_penalty: Mapped[float] = mapped_column(Float, default=0.0)
    data_completeness: Mapped[float] = mapped_column(Float, default=0.0)
    status_band: Mapped[str] = mapped_column(String(80), default="Insufficient Data")
    score_details_json: Mapped[str] = mapped_column(Text, default="{}")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now(), onupdate=lambda: now())

    contacts: Mapped[list["Contact"]] = relationship(cascade="all, delete-orphan")
    issues: Mapped[list["Issue"]] = relationship(cascade="all, delete-orphan")
    milestones: Mapped[list["ConferenceMilestone"]] = relationship(cascade="all, delete-orphan")
    comments_history: Mapped[list["ConferenceComment"]] = relationship(cascade="all, delete-orphan")


class Contact(Base):
    __tablename__ = "contacts"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    role: Mapped[str] = mapped_column(String(80))
    name: Mapped[str] = mapped_column(String(160))
    email: Mapped[str | None] = mapped_column(String(240), nullable=True)
    organization: Mapped[str | None] = mapped_column(String(240), nullable=True)
    phone: Mapped[str | None] = mapped_column(String(80), nullable=True)
    is_primary: Mapped[bool] = mapped_column(Boolean, default=False)
    active: Mapped[bool] = mapped_column(Boolean, default=True)


class Alias(Base):
    __tablename__ = "conference_aliases"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    alias: Mapped[str] = mapped_column(String(240), index=True)
    alias_type: Mapped[str] = mapped_column(String(80), default="Manual")
    source: Mapped[str | None] = mapped_column(String(160), nullable=True)
    active: Mapped[bool] = mapped_column(Boolean, default=True)


class MilestoneDefinition(Base):
    __tablename__ = "milestone_definitions"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    code: Mapped[str] = mapped_column(String(40), unique=True)
    name: Mapped[str] = mapped_column(String(240))
    description: Mapped[str] = mapped_column(Text, default="")
    score_dimension: Mapped[str] = mapped_column(String(120))
    default_weight: Mapped[float] = mapped_column(Float, default=10)
    applicable_sponsorship_types_json: Mapped[str] = mapped_column(Text, default="[]")
    applicable_series_json: Mapped[str] = mapped_column(Text, default="[]")
    required_lifecycle_phase: Mapped[str | None] = mapped_column(String(120), nullable=True)
    due_days_from_start: Mapped[int] = mapped_column(Integer, default=0)
    mandatory: Mapped[bool] = mapped_column(Boolean, default=True)
    enabled: Mapped[bool] = mapped_column(Boolean, default=True)


class ConferenceMilestone(Base):
    __tablename__ = "conference_milestones"
    __table_args__ = (UniqueConstraint("conference_id", "definition_id", name="uq_conf_milestone"),)

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    definition_id: Mapped[str] = mapped_column(ForeignKey("milestone_definitions.id"), index=True)
    status: Mapped[str] = mapped_column(String(80), default="Unknown")
    due_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    completed_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    manual_override: Mapped[bool] = mapped_column(Boolean, default=False)
    comments: Mapped[str | None] = mapped_column(Text, nullable=True)
    source: Mapped[str] = mapped_column(String(80), default="System")
    last_updated: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    definition: Mapped[MilestoneDefinition] = relationship()


class Issue(Base):
    __tablename__ = "issues"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    issue_key: Mapped[str] = mapped_column(String(120), index=True)
    title: Mapped[str] = mapped_column(String(240))
    description: Mapped[str] = mapped_column(Text, default="")
    category: Mapped[str] = mapped_column(String(80), default="Data Quality")
    severity: Mapped[str] = mapped_column(String(40), default="Medium")
    issue_status: Mapped[str] = mapped_column(String(80), default="Open")
    review_assessment: Mapped[str] = mapped_column(String(80), default="Unreviewed")
    source_type: Mapped[str] = mapped_column(String(80), default="Rule")
    source_field: Mapped[str | None] = mapped_column(String(120), nullable=True)
    rule_identifier: Mapped[str | None] = mapped_column(String(120), nullable=True)
    owner: Mapped[str | None] = mapped_column(String(160), nullable=True)
    date_detected: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    due_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    last_reviewed_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    resolution_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    user_comment: Mapped[str | None] = mapped_column(Text, nullable=True)
    ai_recommendation: Mapped[str | None] = mapped_column(Text, nullable=True)
    recommendation_generated_date: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)
    policy_references_json: Mapped[str] = mapped_column(Text, default="[]")
    active: Mapped[bool] = mapped_column(Boolean, default=True)


class IssueComment(Base):
    __tablename__ = "issue_comments"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    issue_id: Mapped[str] = mapped_column(ForeignKey("issues.id"), index=True)
    comment: Mapped[str] = mapped_column(Text)
    author: Mapped[str] = mapped_column(String(160), default="local-user")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())


class ConferenceComment(Base):
    __tablename__ = "conference_comments"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    comment: Mapped[str] = mapped_column(Text)
    author: Mapped[str] = mapped_column(String(160), default="local-user")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now(), onupdate=lambda: now())


class ImportBatch(Base):
    __tablename__ = "import_batches"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    original_filename: Mapped[str] = mapped_column(String(260))
    file_type: Mapped[str] = mapped_column(String(20))
    file_hash: Mapped[str] = mapped_column(String(64), index=True)
    report_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    upload_date: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    import_status: Mapped[str] = mapped_column(String(80), default="Validated")
    rows_count: Mapped[int] = mapped_column(Integer, default=0)
    new_count: Mapped[int] = mapped_column(Integer, default=0)
    changed_count: Mapped[int] = mapped_column(Integer, default=0)
    unchanged_count: Mapped[int] = mapped_column(Integer, default=0)
    conflict_count: Mapped[int] = mapped_column(Integer, default=0)
    notes: Mapped[str | None] = mapped_column(Text, nullable=True)
    preview_json: Mapped[str] = mapped_column(Text, default="{}")
    file_data: Mapped[bytes | None] = mapped_column(LargeBinary, nullable=True)


class FieldChange(Base):
    __tablename__ = "field_change_history"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    entity: Mapped[str] = mapped_column(String(80), default="Conference")
    field_name: Mapped[str] = mapped_column(String(120))
    old_value: Mapped[str | None] = mapped_column(Text, nullable=True)
    new_value: Mapped[str | None] = mapped_column(Text, nullable=True)
    change_type: Mapped[str] = mapped_column(String(80), default="Manual")
    source: Mapped[str] = mapped_column(String(160), default="UI")
    import_batch_id: Mapped[str | None] = mapped_column(ForeignKey("import_batches.id"), nullable=True)
    timestamp: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    comment: Mapped[str | None] = mapped_column(Text, nullable=True)


class Snapshot(Base):
    __tablename__ = "monthly_snapshots"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    snapshot_timestamp: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    payload_json: Mapped[str] = mapped_column(Text)


class ScoreHistory(Base):
    __tablename__ = "score_history"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    score: Mapped[float] = mapped_column(Float)
    data_completeness: Mapped[float] = mapped_column(Float)
    dimension_scores_json: Mapped[str] = mapped_column(Text, default="{}")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())


class Document(Base):
    __tablename__ = "documents"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    title: Mapped[str] = mapped_column(String(260))
    file_name: Mapped[str] = mapped_column(String(260))
    document_category: Mapped[str] = mapped_column(String(120))
    knowledge_scope: Mapped[str] = mapped_column(String(120))
    conference_id: Mapped[str | None] = mapped_column(ForeignKey("conferences.id"), nullable=True)
    conference_series: Mapped[str | None] = mapped_column(String(160), nullable=True)
    version: Mapped[str | None] = mapped_column(String(80), nullable=True)
    effective_date: Mapped[date | None] = mapped_column(Date, nullable=True)
    source_url: Mapped[str | None] = mapped_column(String(500), nullable=True)
    upload_date: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    active: Mapped[bool] = mapped_column(Boolean, default=True)
    extraction_state: Mapped[str] = mapped_column(String(80), default="Extracted")
    indexing_state: Mapped[str] = mapped_column(String(80), default="Indexed")
    page_count: Mapped[int] = mapped_column(Integer, default=1)
    chunk_count: Mapped[int] = mapped_column(Integer, default=0)
    metadata_json: Mapped[str] = mapped_column(Text, default="{}")
    extracted_text: Mapped[str] = mapped_column(Text, default="")


class TemplateFile(Base):
    __tablename__ = "template_files"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    template_name: Mapped[str] = mapped_column(String(260))
    short_description: Mapped[str] = mapped_column(Text, default="")
    category: Mapped[str] = mapped_column(String(120), index=True)
    template_type: Mapped[str] = mapped_column(String(80))
    file_name: Mapped[str] = mapped_column(String(260))
    original_filename: Mapped[str] = mapped_column(String(260))
    file_data: Mapped[bytes | None] = mapped_column(LargeBinary, nullable=True)
    upload_date: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now(), onupdate=lambda: now())


class EmailDraft(Base):
    __tablename__ = "email_drafts"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), index=True)
    related_issues_json: Mapped[str] = mapped_column(Text, default="[]")
    recipient_names: Mapped[str] = mapped_column(Text, default="")
    recipient_addresses: Mapped[str] = mapped_column(Text, default="")
    cc_addresses: Mapped[str] = mapped_column(Text, default="")
    subject: Mapped[str] = mapped_column(String(260))
    body: Mapped[str] = mapped_column(Text)
    tone: Mapped[str] = mapped_column(String(80), default="Concise professional")
    generator: Mapped[str] = mapped_column(String(80), default="Local composer")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=lambda: now(), onupdate=lambda: now())
    saved: Mapped[bool] = mapped_column(Boolean, default=True)


class StatusMapping(Base):
    __tablename__ = "status_mappings"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    source_value: Mapped[str] = mapped_column(String(160), unique=True)
    normalized_value: Mapped[str] = mapped_column(String(80))
    active: Mapped[bool] = mapped_column(Boolean, default=True)


class DashboardPin(Base):
    __tablename__ = "dashboard_pins"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    conference_id: Mapped[str] = mapped_column(ForeignKey("conferences.id"), unique=True)
    display_order: Mapped[int] = mapped_column(Integer, default=0)
    pin_group: Mapped[str] = mapped_column(String(80), default="Default")
    date_pinned: Mapped[datetime] = mapped_column(DateTime, default=lambda: now())


class DashboardSettings(Base):
    __tablename__ = "dashboard_settings"

    key: Mapped[str] = mapped_column(String(120), primary_key=True)
    value_json: Mapped[str] = mapped_column(Text)


@dataclass(frozen=True)
class DashboardState:
    engine: Any
    session_factory: sessionmaker[Session]


_state: DashboardState | None = None


def init_dashboard() -> DashboardState:
    global _state
    load_local_env()
    for directory in [
        app_path("APP_DOCUMENT_PATH", "./data/documents"),
        app_path("APP_IMPORT_PATH", "./data/imports"),
        app_path("APP_EXPORT_PATH", "./data/exports"),
        app_path("APP_VECTOR_PATH", "./data/vector_store"),
        app_path("APP_TEMPLATE_PATH", "./data/templates"),
    ]:
        directory.mkdir(parents=True, exist_ok=True)
    db_path = app_path("APP_DATABASE_PATH", "./data/itss_dashboard.db")
    db_path.parent.mkdir(parents=True, exist_ok=True)
    engine = create_engine(f"sqlite:///{db_path.as_posix()}", future=True)
    Base.metadata.create_all(engine)
    ensure_database_schema(engine)
    factory = sessionmaker(engine, expire_on_commit=False, future=True)
    state = DashboardState(engine=engine, session_factory=factory)
    _state = state
    with factory() as session:
        seed_configuration(session)
        for conference in session.scalars(select(Conference)):
            recalculate(conference, session, record_history=False)
        session.commit()
    return state


def ensure_database_schema(engine: Any) -> None:
    inspector = inspect(engine)
    table_names = inspector.get_table_names()
    if "conferences" not in table_names:
        return
    columns = {column["name"] for column in inspector.get_columns("conferences")}
    additions = {
        "last_source_update": "DATE",
        "source_details_json": "TEXT DEFAULT '{}'",
        "itss_loan_requested": "BOOLEAN DEFAULT 0",
        "itss_loan_amount": "FLOAT",
    }
    with engine.begin() as connection:
        for name, ddl_type in additions.items():
            if name not in columns:
                connection.exec_driver_sql(f"ALTER TABLE conferences ADD COLUMN {name} {ddl_type}")
        if "email_drafts" in table_names:
            email_columns = {column["name"] for column in inspector.get_columns("email_drafts")}
            if "generator" not in email_columns:
                connection.exec_driver_sql("ALTER TABLE email_drafts ADD COLUMN generator VARCHAR(80) DEFAULT 'Local composer'")
        if "template_files" in table_names:
            template_columns = {column["name"] for column in inspector.get_columns("template_files")}
            if "file_data" not in template_columns:
                connection.exec_driver_sql("ALTER TABLE template_files ADD COLUMN file_data BLOB")
        if "import_batches" in table_names:
            import_columns = {column["name"] for column in inspector.get_columns("import_batches")}
            if "file_data" not in import_columns:
                connection.exec_driver_sql("ALTER TABLE import_batches ADD COLUMN file_data BLOB")


def get_state() -> DashboardState:
    return _state or init_dashboard()


def get_session() -> Session:
    state = get_state()
    with state.session_factory() as session:
        yield session


def now() -> datetime:
    return datetime.now(timezone.utc).replace(tzinfo=None)


def normalize_acronym(value: str) -> str:
    return re.sub(r"[^A-Z0-9]", "", value.upper())


def normalize_record_number(value: Any) -> str | None:
    if value is None or pd.isna(value):
        return None
    text = str(value).strip()
    if not text:
        return None
    try:
        numeric = float(text)
        if numeric.is_integer():
            return str(int(numeric))
    except ValueError:
        pass
    return text


def parse_date(value: Any) -> date | None:
    if value is None or value == "" or pd.isna(value):
        return None
    if isinstance(value, date):
        return value
    return date.fromisoformat(str(value)[:10])


def sanitize_reference_config(raw: dict[str, Any] | None = None) -> dict[str, Any]:
    raw = raw or {}
    clean: dict[str, Any] = {}
    for key, defaults in REFERENCE_CONFIG_DEFAULTS.items():
        values = raw.get(key, defaults)
        if key == "conference_series":
            clean[key] = sanitize_conference_series(values)
            continue
        if key == "committee_members":
            clean[key] = sanitize_committee_members(values)
            continue
        if not isinstance(values, list):
            values = defaults
        seen: set[str] = set()
        items: list[str] = []
        for value in values:
            text = " ".join(str(value).strip().split())
            if key == "contact_roles" and text == "Finance Chair or Treasurer":
                text = "Finance Chair"
            if text and text not in seen:
                seen.add(text)
                items.append(text)
        if key not in {"conference_statuses", "review_assessments"} and "Unknown" not in seen:
            items.insert(0, "Unknown")
        if key == "conference_statuses":
            items = [item for item in items if item != "Unknown"]
            for required in DERIVED_CONFERENCE_STATUSES:
                if required not in items:
                    items.append(required)
        if key == "normalized_statuses":
            for required in ("Unknown", "Open", "Resolved", "Closed", "Complete", "Approved"):
                if required not in items:
                    items.append(required)
        if key == "review_assessments":
            items = [item for item in items if item not in {"Unknown", "Acknowledged"}]
            for required in REVIEW_ASSESSMENTS:
                if required not in items:
                    items.append(required)
        clean[key] = items
    return clean


def sanitize_committee_members(values: Any) -> list[str]:
    if not isinstance(values, list):
        return ["Ahmed Hussein", "Daniel Medina"]
    seen: set[str] = set()
    items: list[str] = []
    for value in values:
        text = " ".join(str(value).strip().split())
        if text and text not in seen and text.lower() != "unknown":
            seen.add(text)
            items.append(text)
    return items if items else ["Ahmed Hussein", "Daniel Medina"]


def sanitize_conference_series(values: Any) -> list[dict[str, Any]]:
    if not isinstance(values, list):
        values = REFERENCE_CONFIG_DEFAULTS["conference_series"]
    defaults_by_name = {name.upper(): {"code": code, "name": name, "flagship": flagship} for code, name, flagship in CONFERENCE_SERIES}
    defaults_by_code = {str(code).upper(): {"code": code, "name": name, "flagship": flagship} for code, name, flagship in CONFERENCE_SERIES}
    items: list[dict[str, Any]] = []
    seen: set[str] = set()
    for value in values:
        if isinstance(value, str) and value.strip().startswith("{"):
            try:
                parsed = ast.literal_eval(value.strip())
                if isinstance(parsed, dict):
                    value = parsed
            except (SyntaxError, ValueError):
                pass
        if isinstance(value, dict):
            raw_code = str(value.get("code", "")).strip()
            raw_name = str(value.get("name", "")).strip()
            flagship = bool(value.get("flagship", False))
        else:
            text = " ".join(str(value).strip().split())
            if " - " in text:
                raw_code, raw_name = text.split(" - ", 1)
                flagship = raw_code.strip().upper() in {"ITSC", "IV"}
            else:
                matched = defaults_by_name.get(text.upper()) or defaults_by_code.get(text.upper())
                if matched:
                    raw_code, raw_name = matched["code"], matched["name"]
                    flagship = bool(matched["flagship"])
                else:
                    raw_code, raw_name = text, text
                    flagship = raw_code.strip().upper() in {"ITSC", "IV"}
        matched = defaults_by_name.get(raw_name.upper()) or defaults_by_code.get(raw_code.upper())
        if matched:
            raw_code, raw_name, flagship = matched["code"], matched["name"], bool(matched["flagship"])
        code = raw_code.strip().upper() or "UNKNOWN"
        name = raw_name.strip() or code
        if code in seen:
            continue
        seen.add(code)
        items.append({"code": code, "name": name, "flagship": flagship})
    if "UNKNOWN" not in seen:
        items.insert(0, {"code": "UNKNOWN", "name": "Unknown", "flagship": False})
    return items


def reference_config(session: Session) -> dict[str, Any]:
    setting = session.get(DashboardSettings, "reference_config")
    loaded: dict[str, Any] = {}
    if setting:
        try:
            loaded = json.loads(setting.value_json)
        except json.JSONDecodeError:
            loaded = {}
    clean = sanitize_reference_config(loaded)
    if not setting:
        session.add(DashboardSettings(key="reference_config", value_json=json.dumps(clean)))
    elif loaded != clean:
        setting.value_json = json.dumps(clean)
    return clean


def allowed_reference_values(session: Session, key: str) -> list[str]:
    values = reference_config(session).get(key, [])
    if key == "conference_series":
        allowed: list[str] = []
        for item in values:
            if isinstance(item, dict):
                allowed.extend([str(item.get("code", "")), str(item.get("name", ""))])
        return [value for value in allowed if value]
    return values


def normalize_status(raw: str | None, session: Session | None = None) -> str:
    if not raw:
        return "Unknown"
    cleaned = " ".join(str(raw).strip().split())
    allowed_statuses = allowed_reference_values(session, "normalized_statuses") if session is not None else NORMALIZED_STATUSES
    if cleaned in allowed_statuses:
        return cleaned
    if session is not None:
        mapping = session.scalar(select(StatusMapping).where(StatusMapping.source_value == cleaned))
        if mapping and mapping.active and mapping.normalized_value in allowed_statuses:
            return mapping.normalized_value
    lowered = cleaned.lower()
    if "complete" in lowered or "done" in lowered:
        return "Complete" if "Complete" in allowed_statuses else "Unknown"
    if "approve" in lowered:
        return "Approved" if "Approved" in allowed_statuses else "Unknown"
    if "submit" in lowered:
        return "Submitted" if "Submitted" in allowed_statuses else "Unknown"
    if "publish" in lowered or "xplore" in lowered:
        return "Published" if "Published" in allowed_statuses else "Unknown"
    if "block" in lowered or "hold" in lowered:
        return "Blocked" if "Blocked" in allowed_statuses else "Unknown"
    if "cancel" in lowered:
        return "Cancelled" if "Cancelled" in allowed_statuses else "Unknown"
    if "progress" in lowered:
        return "In Progress" if "In Progress" in allowed_statuses else "Unknown"
    if "not" in lowered and "start" in lowered:
        return "Not Started" if "Not Started" in allowed_statuses else "Unknown"
    return "Unknown"


def seed_configuration(session: Session) -> None:
    if session.scalar(select(func.count(MilestoneDefinition.id))) == 0:
        for code, name, dimension, _target, due_days in MILESTONE_SEEDS:
            session.add(
                MilestoneDefinition(
                    code=code,
                    name=name,
                    description=f"Default readiness milestone for {dimension}.",
                    score_dimension=dimension,
                    default_weight=SCORE_WEIGHTS[dimension],
                    due_days_from_start=due_days,
                    mandatory=True,
                    enabled=True,
                )
            )
    for status in REFERENCE_CONFIG_DEFAULTS["normalized_statuses"]:
        if not session.scalar(select(StatusMapping).where(StatusMapping.source_value == status)):
            session.add(StatusMapping(source_value=status, normalized_value=status))
    if not session.get(DashboardSettings, "score_weights"):
        session.add(DashboardSettings(key="score_weights", value_json=json.dumps(SCORE_WEIGHTS)))
    if not session.get(DashboardSettings, "score_settings"):
        session.add(DashboardSettings(key="score_settings", value_json=json.dumps(DEFAULT_SCORE_SETTINGS)))
    if not session.get(DashboardSettings, "portfolio_start_year"):
        session.add(DashboardSettings(key="portfolio_start_year", value_json="2020"))
    if not session.get(DashboardSettings, "feature_flags"):
        session.add(DashboardSettings(key="feature_flags", value_json=json.dumps(default_feature_flags())))
    if not session.get(DashboardSettings, "role_permissions"):
        session.add(DashboardSettings(key="role_permissions", value_json=json.dumps(default_role_permissions())))
    if not session.get(DashboardSettings, "assistant_system_prompt"):
        session.add(DashboardSettings(key="assistant_system_prompt", value_json=json.dumps(DEFAULT_ASSISTANT_SYSTEM_PROMPT)))
    reference_config(session)


def default_feature_flags() -> dict[str, bool]:
    return {
        "conferences": True,
        "issues": True,
        "imports": True,
        "knowledge_base": True,
        "templates": True,
        "assistant": True,
        "email_drafts": True,
        "exports": True,
        "system_status": True,
    }


def merged_score_settings(raw: dict[str, Any] | None = None) -> dict[str, Any]:
    raw = raw if isinstance(raw, dict) else {}
    clean = json.loads(json.dumps(DEFAULT_SCORE_SETTINGS))
    for key in ("dimension_weights", "milestone_status_scores", "issue_severity_penalties", "issue_assessment_factors"):
        incoming = raw.get(key, {})
        if isinstance(incoming, dict):
            for item_key, value in incoming.items():
                try:
                    clean[key][str(item_key)] = float(value)
                except (TypeError, ValueError):
                    continue
    try:
        clean["issue_penalty_cap"] = max(0.0, float(raw.get("issue_penalty_cap", clean["issue_penalty_cap"])))
    except (TypeError, ValueError):
        pass
    for key in ("lateness_step_days", "lateness_cap_factor"):
        try:
            clean[key] = max(0.0, float(raw.get(key, clean[key])))
        except (TypeError, ValueError):
            pass
    formula = str(raw.get("score_formula", clean["score_formula"])).strip()
    if formula:
        clean["score_formula"] = formula
    return clean


SCORE_FORMULA_FUNCTIONS = {"min": min, "max": max, "round": round, "abs": abs}
SCORE_FORMULA_ALLOWED_NODES = (
    ast.Expression,
    ast.BinOp,
    ast.UnaryOp,
    ast.Call,
    ast.Name,
    ast.Load,
    ast.Constant,
    ast.Add,
    ast.Sub,
    ast.Mult,
    ast.Div,
    ast.Pow,
    ast.Mod,
    ast.USub,
    ast.UAdd,
)


def evaluate_score_formula(formula: str, context: dict[str, float]) -> float:
    try:
        tree = ast.parse(formula, mode="eval")
    except SyntaxError as exc:
        raise ValueError(f"Invalid score formula syntax: {exc.msg}") from exc
    allowed_names = set(context) | set(SCORE_FORMULA_FUNCTIONS)
    for node in ast.walk(tree):
        if not isinstance(node, SCORE_FORMULA_ALLOWED_NODES):
            raise ValueError(f"Unsupported score formula element: {type(node).__name__}")
        if isinstance(node, ast.Name) and node.id not in allowed_names:
            raise ValueError(f"Unknown score formula variable: {node.id}")
        if isinstance(node, ast.Call):
            if not isinstance(node.func, ast.Name) or node.func.id not in SCORE_FORMULA_FUNCTIONS:
                raise ValueError("Only min, max, round, and abs can be used as score formula functions.")
    result = eval(
        compile(tree, "<score_formula>", "eval"),
        {"__builtins__": {}},
        {**SCORE_FORMULA_FUNCTIONS, **context},
    )
    try:
        return float(result)
    except (TypeError, ValueError) as exc:
        raise ValueError("Score formula must return a number.") from exc


def validate_score_formula(formula: str) -> None:
    evaluate_score_formula(
        formula,
        {
            "base_score": 78.0,
            "issue_penalty": 6.0,
            "data_completeness": 82.0,
            "milestone_completion_pct": 65.0,
            "total_milestones": 10.0,
            "completed_milestones": 6.0,
            "overdue_milestones": 1.0,
            "blocked_milestones": 0.0,
            "active_milestones": 3.0,
            "due_soon_milestones": 2.0,
        },
    )


def score_settings(session: Session) -> dict[str, Any]:
    setting = session.get(DashboardSettings, "score_settings")
    loaded: dict[str, Any] = {}
    if setting:
        try:
            loaded = json.loads(setting.value_json)
        except json.JSONDecodeError:
            loaded = {}
    clean = merged_score_settings(loaded)
    if not setting:
        session.add(DashboardSettings(key="score_settings", value_json=json.dumps(clean)))
    elif loaded != clean:
        setting.value_json = json.dumps(clean)
    return clean


def default_role_permissions() -> dict[str, dict[str, bool]]:
    return {
        "administrator": {
            "overview": True,
            "conferences": True,
            "assigned_conferences_only": False,
            "conference_edit": True,
            "issues": True,
            "issue_edit": True,
            "imports": True,
            "knowledge_base": True,
            "templates": True,
            "template_upload": True,
            "assistant": True,
            "email_drafts": True,
            "system_status": True,
            "settings": True,
        },
        "conference_organizer": {
            "overview": False,
            "conferences": True,
            "assigned_conferences_only": True,
            "conference_edit": True,
            "issues": False,
            "issue_edit": True,
            "imports": False,
            "knowledge_base": False,
            "templates": True,
            "template_upload": False,
            "assistant": True,
            "email_drafts": False,
            "system_status": True,
            "settings": False,
        },
        "itss_leadership": {
            "overview": True,
            "conferences": False,
            "assigned_conferences_only": False,
            "conference_edit": False,
            "issues": False,
            "issue_edit": False,
            "imports": False,
            "knowledge_base": False,
            "templates": True,
            "template_upload": False,
            "assistant": True,
            "email_drafts": False,
            "system_status": True,
            "settings": False,
        },
        "cee_staff": {
            "overview": True,
            "conferences": True,
            "assigned_conferences_only": False,
            "conference_edit": False,
            "issues": True,
            "issue_edit": False,
            "imports": False,
            "knowledge_base": False,
            "templates": True,
            "template_upload": False,
            "assistant": True,
            "email_drafts": False,
            "system_status": True,
            "settings": False,
        },
    }


def role_catalog() -> list[dict[str, str]]:
    return [
        {"key": "administrator", "label": "Administrator", "description": "Access everything."},
        {"key": "conference_organizer", "label": "Conference Organizer", "description": "Assigned conferences, assistant, and system status."},
        {"key": "itss_leadership", "label": "ITSS Leadership", "description": "Overview, assistant, and system status."},
        {"key": "cee_staff", "label": "CEE Staff", "description": "Overview and all conferences in read-only mode, plus assistant and status."},
    ]


def permission_catalog() -> list[dict[str, str]]:
    return [
        {"key": "overview", "label": "Overview", "description": "Dashboard overview page."},
        {"key": "conferences", "label": "Conferences", "description": "Conference table and detail pages."},
        {"key": "assigned_conferences_only", "label": "Assigned conferences only", "description": "Restrict conference scope once assignments exist."},
        {"key": "conference_edit", "label": "Edit conferences", "description": "Create conferences, edit facts, milestones, contacts, and comments."},
        {"key": "issues", "label": "Issues", "description": "Issue board and issue lists."},
        {"key": "issue_edit", "label": "Edit issues", "description": "Create, edit, resolve, delete, and request AI issue recommendations."},
        {"key": "imports", "label": "Import and export", "description": "Import approval center and exports."},
        {"key": "knowledge_base", "label": "Knowledge base", "description": "Document index and knowledge uploads."},
        {"key": "templates", "label": "Templates", "description": "Download and manage conference organizer template files."},
        {"key": "template_upload", "label": "Manage templates", "description": "Upload and delete template files."},
        {"key": "assistant", "label": "Assistant", "description": "Ask questions through the assistant."},
        {"key": "email_drafts", "label": "Email drafts", "description": "Generate and manage draft emails."},
        {"key": "system_status", "label": "System status", "description": "Service, database, and LLM status pages."},
        {"key": "settings", "label": "Settings", "description": "Portal configuration and role access settings."},
    ]


def merged_role_permissions(raw: dict[str, Any] | None) -> dict[str, dict[str, bool]]:
    defaults = default_role_permissions()
    clean: dict[str, dict[str, bool]] = {}
    for role, permissions in defaults.items():
        incoming = raw.get(role, {}) if isinstance(raw, dict) else {}
        clean[role] = {key: bool(incoming.get(key, value)) for key, value in permissions.items()}
    return clean


def milestone_date_offsets(session: Session | None = None) -> dict[str, dict[str, int]]:
    """Get persisted milestone date offsets, falling back to MILESTONE_DATE_DEFAULTS."""
    def normalize_offsets(raw: dict[str, Any]) -> dict[str, dict[str, Any]]:
        merged = json.loads(json.dumps(MILESTONE_DATE_DEFAULTS))
        for code, offset in raw.items():
            if not isinstance(offset, dict):
                continue
            current = dict(merged.get(str(code).upper(), {"anchor": "start", "months": 0, "days": 0}))
            anchor = str(offset.get("anchor", current.get("anchor", "start"))).lower()
            current["anchor"] = anchor if anchor in {"start", "end"} else "start"
            try:
                current["months"] = int(offset.get("months", current.get("months", 0)))
            except (TypeError, ValueError):
                current["months"] = int(current.get("months", 0))
            try:
                current["days"] = int(offset.get("days", current.get("days", 0)))
            except (TypeError, ValueError):
                current["days"] = int(current.get("days", 0))
            merged[str(code).upper()] = current
        return merged

    if session is None:
        return normalize_offsets(MILESTONE_DATE_DEFAULTS)
    setting = session.get(DashboardSettings, "milestone_date_defaults")
    if setting:
        try:
            loaded = json.loads(setting.value_json)
            if isinstance(loaded, dict):
                return normalize_offsets(loaded)
        except (json.JSONDecodeError, TypeError):
            pass
    return normalize_offsets(MILESTONE_DATE_DEFAULTS)


def milestone_due_date(code: str, start_date: date | None, end_date: date | None, *, offsets: dict[str, dict[str, Any]] | None = None) -> date | None:
    """Calculate milestone due date from conference start/end date using months/days offset."""
    lookup = offsets if offsets is not None else MILESTONE_DATE_DEFAULTS
    offset = lookup.get(code)
    if offset is None:
        return None
    anchor = str(offset.get("anchor", "start")).lower()
    anchor_date = end_date if anchor == "end" else start_date
    if anchor_date is None:
        anchor_date = start_date or end_date
    if anchor_date is None:
        return None
    total_days = offset["months"] * 30 + offset["days"]
    from datetime import timedelta
    result = anchor_date + timedelta(days=total_days)
    return result


def milestone_due_from_start(code: str, start_date: date | None, *, offsets: dict[str, dict[str, int]] | None = None) -> date | None:
    return milestone_due_date(code, start_date, None, offsets=offsets)


def ensure_milestones(conference: Conference, session: Session) -> None:
    definitions = list(session.scalars(select(MilestoneDefinition).where(MilestoneDefinition.enabled.is_(True))))
    existing = {item.definition_id for item in conference.milestones}
    offsets = milestone_date_offsets(session)
    for definition in definitions:
        if definition.id in existing:
            continue
        due = None
        if conference.start_date:
            due = milestone_due_date(definition.code, conference.start_date, conference.end_date, offsets=offsets)
        status = "Unknown"
        if definition.code == "APPLICATION":
            status = conference.application_status
        elif definition.code == "MOU":
            status = conference.mou_status
        elif definition.code in {"BUDGET", "BANKING"}:
            status = conference.finance_status
        elif definition.code in {"PROCEEDINGS", "REVIEWS"}:
            status = conference.publication_status
        session.add(
            ConferenceMilestone(
                conference_id=conference.id,
                definition_id=definition.id,
                status=status,
                due_date=due,
            )
        )


def suggest_phase(conference: Conference) -> str:
    today = date.today()
    if conference.conference_status in {"Closed", "Cancelled"}:
        return conference.conference_status
    if conference.end_date and conference.end_date < today:
        if conference.xplore_posting_date and conference.accounting_close_date:
            return "Closed"
        if conference.proceedings_submitted_date or conference.publication_status in {"Submitted", "Published"}:
            return "Proceedings Processing"
        return "Financial and Administrative Closure"
    if conference.start_date:
        days = (conference.start_date - today).days
        if days <= 7:
            return "Conference Delivery"
        if days <= 90:
            return "Registration and Final Preparation"
        if conference.submission_deadline and conference.submission_deadline <= today:
            return "Submission and Review"
        if days <= 365:
            return "Detailed Planning"
    if conference.application_status in {"Approved", "Complete"} and conference.mou_status in {"Approved", "Complete"}:
        return "Detailed Planning"
    if conference.application_status in {"Submitted", "Awaiting IEEE", "In Progress"}:
        return "IEEE Application and MOU"
    if conference.sponsorship_type:
        return "ITSS Approved"
    return "Expression of Interest"


def lateness_factor(due_date: date | None, step_days: float = 30.0, cap_factor: float = 3.0) -> float:
    """Calculate lateness multiplier for a milestone.

    Returns 0 if the due date is in the future or today.
    For overdue milestones, returns a linear ramp:
      overdue_days / step_days, capped at cap_factor.

    This factor multiplies the milestone weight so that overdue
    milestones count more heavily against the final score.
    """
    if due_date is None:
        return 0.0
    days_past = (due_date - date.today()).days
    if days_past >= 0:
        return 0.0
    overdue_days = -days_past
    return min(overdue_days / step_days, cap_factor)


def milestone_score(status: str, due_date: date | None, status_scores: dict[str, float] | None = None) -> float | None:
    scores = {**DEFAULT_MILESTONE_STATUS_SCORES, **(status_scores or {})}
    status = normalize_status(status)
    if status in {"Cancelled", "Not Applicable"}:
        return None
    if status in {"Approved", "Complete", "Published", "Closed"}:
        return float(scores["completed"])
    if status == "Unknown":
        return float(scores["unknown"])
    if due_date is None:
        return float(scores["no_due_date"])
    days = (due_date - date.today()).days
    if status == "Not Started":
        if days > 90:
            return float(scores["not_started_far"])
        if days > 30:
            return float(scores["not_started_upcoming"])
        if days > 0:
            return float(scores["not_started_due_soon"])
        return float(scores["not_started_overdue"])
    if status == "In Progress":
        if days >= 0:
            return float(scores["in_progress_on_time"])
        if days >= -30:
            return float(scores["in_progress_recently_overdue"])
        return float(scores["in_progress_overdue"])
    if status in {"Submitted", "Awaiting IEEE", "Awaiting Conference", "Awaiting External Party"}:
        if days >= 0:
            return float(scores["awaiting_on_time"])
        if days >= -30:
            return float(scores["awaiting_recently_overdue"])
        return float(scores["awaiting_overdue"])
    if status in {"Blocked", "Rejected"}:
        return float(scores["blocked"])
    return None


def issue_penalty(issues: list[Issue], settings: dict[str, Any]) -> float:
    severity_points = settings.get("issue_severity_penalties", DEFAULT_SCORE_SETTINGS["issue_severity_penalties"])
    assessment_factor = settings.get("issue_assessment_factors", DEFAULT_SCORE_SETTINGS["issue_assessment_factors"])
    cap = float(settings.get("issue_penalty_cap", DEFAULT_SCORE_SETTINGS["issue_penalty_cap"]))
    total = sum(
        severity_points.get(issue.severity, 2.0) * assessment_factor.get(issue.review_assessment, 1.0)
        for issue in issues
        if issue.active and issue.issue_status != "Resolved"
    )
    return min(cap, total)


def data_completeness(conference: Conference) -> float:
    fields = [
        conference.acronym,
        conference.year,
        conference.official_title,
        conference.conference_series,
        conference.sponsorship_type,
        conference.lifecycle_phase,
        conference.start_date,
        conference.end_date,
        conference.city,
        conference.country,
        conference.website,
        conference.estimated_attendees,
        conference.application_status,
        conference.mou_status,
        conference.finance_status,
        conference.publication_status,
    ]
    milestone_values = [m.status for m in conference.milestones]
    contact_complete = any(contact.email for contact in conference.contacts if contact.active)
    filled = sum(1 for item in fields if item not in (None, "", "Unknown"))
    filled += sum(1 for status in milestone_values if status != "Unknown")
    filled += 1 if contact_complete else 0
    total = len(fields) + max(len(milestone_values), 1) + 1
    return round((filled / total) * 100, 1)


def status_band(score: float, completeness: float, conference_status: str) -> str:
    if conference_status == "Cancelled":
        return "Cancelled"
    if conference_status == "Closed":
        return "Closed"
    if conference_status in {"Blocked", "Critical", "At Risk", "Attention Needed", "On Track"}:
        return conference_status
    if completeness < 35:
        return "Provisional"
    if score >= 85:
        return "On Track"
    if score >= 70:
        return "Attention Needed"
    if score >= 50:
        return "At Risk"
    return "Critical"


def sync_conference_facts_from_milestones(conference: Conference) -> None:
    statuses = {milestone.definition.code: milestone.status for milestone in conference.milestones}
    if statuses.get("APPLICATION"):
        conference.application_status = statuses["APPLICATION"]
    if statuses.get("MOU"):
        conference.mou_status = statuses["MOU"]
    finance_statuses = [statuses[code] for code in ("BUDGET", "BANKING", "FIN_CLOSE") if code in statuses]
    if finance_statuses:
        conference.finance_status = aggregate_finance_status(statuses)
    publication_statuses = [statuses[code] for code in ("CFP", "REVIEWS", "PROCEEDINGS") if code in statuses]
    if publication_statuses:
        conference.publication_status = aggregate_status(publication_statuses)


COMPLETE_STATUSES = {"Approved", "Complete", "Published", "Closed", "Not Applicable"}
ACTIVE_STATUSES = {"In Progress", "Submitted", "Awaiting IEEE", "Awaiting Conference", "Awaiting External Party"}
BLOCKED_STATUSES = {"Blocked", "Rejected"}
IGNORED_STATUSES = {"Cancelled", "Not Applicable"}


def milestone_is_complete(milestone: ConferenceMilestone) -> bool:
    return normalize_status(milestone.status) in COMPLETE_STATUSES


def milestone_is_actionable(milestone: ConferenceMilestone) -> bool:
    return normalize_status(milestone.status) not in IGNORED_STATUSES


def milestone_stats(conference: Conference) -> dict[str, Any]:
    today = date.today()
    milestones = [m for m in conference.milestones if milestone_is_actionable(m)]
    total = len(milestones)
    completed = [m for m in milestones if milestone_is_complete(m)]
    active = [m for m in milestones if normalize_status(m.status) in ACTIVE_STATUSES]
    blocked = [m for m in milestones if normalize_status(m.status) in BLOCKED_STATUSES]
    unfinished = [m for m in milestones if not milestone_is_complete(m)]
    overdue = [m for m in unfinished if m.due_date and m.due_date < today]
    due_soon = [m for m in unfinished if m.due_date and today <= m.due_date <= today + timedelta(days=30)]
    max_overdue_days = max(((today - m.due_date).days for m in overdue if m.due_date), default=0)
    completion_pct = (len(completed) / total) * 100 if total else 0.0
    return {
        "total": total,
        "completed": completed,
        "active": active,
        "blocked": blocked,
        "unfinished": unfinished,
        "overdue": overdue,
        "due_soon": due_soon,
        "max_overdue_days": max_overdue_days,
        "completion_pct": completion_pct,
    }


def milestone_code_status(conference: Conference, code: str) -> str:
    for milestone in conference.milestones:
        if milestone.definition.code == code:
            return normalize_status(milestone.status)
    return "Unknown"


def milestone_code_complete(conference: Conference, code: str) -> bool:
    return milestone_code_status(conference, code) in COMPLETE_STATUSES


def milestone_code_started(conference: Conference, code: str) -> bool:
    return milestone_code_status(conference, code) in COMPLETE_STATUSES | ACTIVE_STATUSES | BLOCKED_STATUSES


def phase_for_milestone_code(code: str) -> str:
    if code in {"APPLICATION", "MOU"}:
        return "IEEE Application and MOU"
    if code in {"BUDGET", "BANKING", "CFP", "VENUE"}:
        return "Detailed Planning"
    if code == "REVIEWS":
        return "Submission and Review"
    if code == "REGISTRATION":
        return "Registration and Final Preparation"
    if code == "PROCEEDINGS":
        return "Proceedings Processing"
    if code == "FIN_CLOSE":
        return "Financial and Administrative Closure"
    return "Detailed Planning"


def derive_lifecycle_phase(conference: Conference) -> str:
    today = date.today()
    stats = milestone_stats(conference)
    unfinished = sorted(
        stats["unfinished"],
        key=lambda item: (item.due_date or date.max, item.definition.due_days_from_start, item.definition.code),
    )
    if conference.end_date and conference.end_date < today:
        if milestone_code_complete(conference, "FIN_CLOSE") and milestone_code_complete(conference, "PROCEEDINGS"):
            return "Closed"
        if not milestone_code_complete(conference, "PROCEEDINGS"):
            return "Proceedings Processing"
        return "Financial and Administrative Closure"
    if conference.start_date and conference.end_date and conference.start_date <= today <= conference.end_date:
        return "Conference Delivery"
    if conference.start_date and 0 <= (conference.start_date - today).days <= 30:
        return "Registration and Final Preparation"
    if unfinished:
        return phase_for_milestone_code(unfinished[0].definition.code)
    if stats["total"] and stats["completion_pct"] >= 100:
        return "Closed" if conference.end_date and conference.end_date < today else "Conference Delivery"
    if milestone_code_started(conference, "APPLICATION") or milestone_code_started(conference, "MOU"):
        return "IEEE Application and MOU"
    return "Expression of Interest"


def score_status(score: float) -> str:
    if score < 50:
        return "Critical"
    if score < 70:
        return "At Risk"
    if score < 85:
        return "Attention Needed"
    return "On Track"


def derive_conference_status(conference: Conference, score: float, lifecycle_phase: str) -> str:
    stats = milestone_stats(conference)
    if conference.conference_status == "Cancelled":
        return "Cancelled"
    if lifecycle_phase == "Closed":
        return "Closed"
    if stats["total"] and stats["completion_pct"] >= 100:
        return "Complete"
    return score_status(score)


def aggregate_status(statuses: list[str]) -> str:
    values = [status for status in statuses if status]
    known_values = [status for status in values if status != "Unknown"]
    if not known_values:
        return "Unknown"
    priority = [
        "Blocked",
        "Rejected",
        "Not Started",
        "In Progress",
        "Awaiting IEEE",
        "Awaiting Conference",
        "Awaiting External Party",
        "Submitted",
        "Approved",
        "Complete",
        "Published",
        "Closed",
        "Not Applicable",
        "Cancelled",
    ]
    for status in priority:
        if status in known_values:
            return status
    return known_values[0]


def aggregate_finance_status(statuses: dict[str, str]) -> str:
    budget = statuses.get("BUDGET", "Unknown")
    banking = statuses.get("BANKING", "Unknown")
    close = statuses.get("FIN_CLOSE", "Unknown")
    finance_statuses = [budget, banking, close]
    if any(status in {"Blocked", "Rejected"} for status in finance_statuses):
        return aggregate_status(finance_statuses)
    if close in {"Closed", "Complete"}:
        return close
    if budget in {"Approved", "Complete"} and banking in {"Approved", "Complete"}:
        return "In Progress"
    return aggregate_status(finance_statuses)


def recalculate(conference: Conference, session: Session, *, record_history: bool = True) -> None:
    settings = score_settings(session)
    weights = settings.get("dimension_weights", SCORE_WEIGHTS)
    milestone_status_scores = settings.get("milestone_status_scores", DEFAULT_MILESTONE_STATUS_SCORES)
    ensure_milestones(conference, session)
    session.flush()
    sync_conference_facts_from_milestones(conference)
    dimension_totals: dict[str, float] = {}
    dimension_weights: dict[str, float] = {}
    step_days = float(settings.get("lateness_step_days", 30.0))
    cap_factor = float(settings.get("lateness_cap_factor", 3.0))
    weighted = 0.0
    total_weight = 0.0
    details: list[dict[str, Any]] = []
    for milestone in conference.milestones:
        definition = milestone.definition
        score = milestone_score(milestone.status, milestone.due_date, milestone_status_scores)
        if score is None:
            continue
        base_weight = float(weights.get(definition.score_dimension, definition.default_weight))
        lf = lateness_factor(milestone.due_date, step_days, cap_factor)
        effective_weight = base_weight * (1.0 + lf)
        weighted += effective_weight * score
        total_weight += effective_weight
        dimension_totals[definition.score_dimension] = dimension_totals.get(definition.score_dimension, 0) + effective_weight * score
        dimension_weights[definition.score_dimension] = dimension_weights.get(definition.score_dimension, 0) + effective_weight
        details.append({"code": definition.code, "name": definition.name, "status": milestone.status, "score": score, "weight": base_weight, "lateness_factor": round(lf, 3), "effective_weight": round(effective_weight, 1)})
    base = round(weighted / total_weight, 1) if total_weight else 0.0
    penalty = issue_penalty(conference.issues, settings)
    completeness = data_completeness(conference)
    stats = milestone_stats(conference)
    formula = str(settings.get("score_formula", DEFAULT_SCORE_SETTINGS["score_formula"])).strip() or DEFAULT_SCORE_SETTINGS["score_formula"]
    formula_context = {
        "base_score": base,
        "issue_penalty": penalty,
        "data_completeness": completeness,
        "milestone_completion_pct": round(float(stats["completion_pct"]), 1),
        "total_milestones": float(stats["total"]),
        "completed_milestones": float(len(stats["completed"])),
        "overdue_milestones": float(len(stats["overdue"])),
        "blocked_milestones": float(len(stats["blocked"])),
        "active_milestones": float(len(stats["active"])),
        "due_soon_milestones": float(len(stats["due_soon"])),
    }
    try:
        raw_score = evaluate_score_formula(formula, formula_context)
    except ValueError:
        formula = DEFAULT_SCORE_SETTINGS["score_formula"]
        raw_score = evaluate_score_formula(formula, formula_context)
    score = round(max(0.0, min(100.0, raw_score)), 1)
    dimension_scores = {
        name: round(dimension_totals[name] / dimension_weights[name], 1)
        for name in dimension_totals
        if dimension_weights.get(name)
    }
    derived_phase = derive_lifecycle_phase(conference)
    allowed_phases = allowed_reference_values(session, "lifecycle_phases")
    if derived_phase not in allowed_phases:
        derived_phase = "Unknown"
    derived_status = derive_conference_status(conference, score, derived_phase)
    allowed_statuses = allowed_reference_values(session, "conference_statuses")
    if derived_status not in allowed_statuses:
        derived_status = "Unknown"
    conference.base_score = base
    conference.issue_penalty = penalty
    conference.score = score
    conference.data_completeness = completeness
    conference.suggested_phase = derived_phase
    conference.lifecycle_phase = derived_phase
    conference.phase_override = False
    conference.conference_status = derived_status
    conference.status_band = status_band(score, completeness, derived_status)
    conference.score_details_json = json.dumps(
        {
            "milestones": details,
            "dimension_scores": dimension_scores,
            "formula": formula,
            "formula_context": formula_context,
        }
    )
    if record_history:
        session.add(
            ScoreHistory(
                conference_id=conference.id,
                score=score,
                data_completeness=completeness,
                dimension_scores_json=json.dumps(dimension_scores),
            )
        )


def detect_issues(conference: Conference, session: Session) -> None:
    return
    rules = [
        ("MISSING_RECORD", "IEEE Conference Record Number is missing", "Governance", "Low", not conference.conference_number),
        ("MISSING_DATES", "Start or end date is missing", "Operations", "Medium", not conference.start_date or not conference.end_date),
        ("MISSING_CONTACT", "No active contact email is available", "Data Quality", "Medium", not any(c.email for c in conference.contacts if c.active)),
        ("MOU_OPEN", "MOU is not approved or complete", "Governance", "High", conference.mou_status not in {"Approved", "Complete", "Not Applicable"}),
        ("FINANCE_UNKNOWN", "Finance status is unknown", "Finance", "High", conference.finance_status == "Unknown"),
        ("PUBLICATION_UNKNOWN", "Publication status is unknown", "Publication", "Medium", conference.publication_status == "Unknown"),
    ]
    existing = {
        issue.rule_identifier: issue
        for issue in session.scalars(select(Issue).where(Issue.conference_id == conference.id, Issue.source_type == "Rule"))
    }
    for key, title, category, severity, triggered in rules:
        issue = existing.get(key)
        if triggered and issue is None:
            session.add(
                Issue(
                    conference_id=conference.id,
                    issue_key=f"{conference.canonical_name}-{key}",
                    title=title,
                    description=f"{title} for {conference.canonical_name}.",
                    category=category,
                    severity=severity,
                    rule_identifier=key,
                )
            )
        elif triggered and issue is not None:
            issue.active = True
            issue.issue_status = "Open" if issue.issue_status == "Resolved" else issue.issue_status
        elif issue is not None and issue.active:
            issue.review_assessment = "Not an Issue"
            issue.issue_status = "Resolved"
            issue.resolution_date = date.today()
            issue.active = False


class ConferenceIn(BaseModel):
    acronym: str = Field(min_length=1, max_length=24)
    year: int = Field(ge=1900, le=2200)
    official_title: str = Field(min_length=1, max_length=500)
    conference_series: str
    sponsorship_type: str
    lifecycle_phase: str
    conference_number: str | None = None
    parent_conference_id: str | None = None
    conference_category: str = "Portfolio"
    conference_status: str = "Unknown"
    start_date: date | None = None
    end_date: date | None = None
    submission_deadline: date | None = None
    notification_date: date | None = None
    camera_ready_deadline: date | None = None
    city: str | None = None
    state_province: str | None = None
    country: str | None = None
    ieee_region: str | None = None
    venue: str | None = None
    website: str | None = None
    estimated_attendees: int | None = None
    primary_contact: str | None = None
    primary_contact_email: EmailStr | None = None

    @field_validator("lifecycle_phase")
    @classmethod
    def valid_phase(cls, value: str) -> str:
        if not value.strip():
            raise ValueError("Unknown lifecycle phase.")
        return value


class ConferenceUpdate(BaseModel):
    official_title: str | None = None
    conference_number: str | None = None
    lifecycle_phase: str | None = None
    phase_override: bool | None = None
    conference_status: str | None = None
    start_date: date | None = None
    end_date: date | None = None
    city: str | None = None
    country: str | None = None
    website: str | None = None
    estimated_attendees: int | None = None
    actual_attendees: int | None = None
    application_status: str | None = None
    mou_status: str | None = None
    finance_status: str | None = None
    publication_status: str | None = None
    total_income_current: float | None = None
    total_expense_current: float | None = None
    budgeted_income_total: float | None = None
    budgeted_expense_total: float | None = None
    itss_loan_requested: bool | None = None
    itss_loan_amount: float | None = None
    comments: str | None = None
    committee_contact: str | None = None
    change_comment: str | None = None


class ConferenceDeleteIn(BaseModel):
    confirmation_record_number: str


class GeneratedIssue(BaseModel):
    title: str
    description: str = ""
    category: str = "AI Review"
    severity: str = "Medium"
    review_assessment: str = "Needs Follow-up"
    owner: str | None = None
    due_date: date | None = None


class IssueIn(BaseModel):
    conference_id: str
    title: str
    description: str = ""
    category: str = "Data Quality"
    severity: str = "Medium"
    owner: str | None = None
    due_date: date | None = None


class IssueUpdate(BaseModel):
    title: str | None = None
    description: str | None = None
    category: str | None = None
    severity: str | None = None
    issue_status: str | None = None
    review_assessment: str | None = None
    owner: str | None = None
    due_date: date | None = None
    user_comment: str | None = None


class ContactIn(BaseModel):
    role: str
    name: str = Field(min_length=1, max_length=160)
    email: EmailStr | None = None
    organization: str | None = None
    phone: str | None = None
    is_primary: bool = False


class ContactUpdate(BaseModel):
    role: str | None = None
    name: str | None = None
    email: EmailStr | None = None
    organization: str | None = None
    phone: str | None = None
    is_primary: bool | None = None
    active: bool | None = None


class ConferenceCommentIn(BaseModel):
    comment: str = Field(min_length=0, max_length=20000)


class MilestoneUpdate(BaseModel):
    status: str
    due_date: date | None = None
    comments: str | None = None

    @field_validator("status")
    @classmethod
    def valid_status(cls, value: str) -> str:
        if not value.strip():
            raise ValueError("Unknown milestone status.")
        return value


class EmailDraftIn(BaseModel):
    conference_id: str
    issue_ids: list[str] = []
    recipient_names: list[str] = []
    recipient_addresses: list[EmailStr] = []
    cc_addresses: list[EmailStr] = []
    purpose: str = "Request status update"
    instructions: str | None = None
    tone: str = "Concise, friendly, professional, direct, action-oriented"


class ChatIn(BaseModel):
    question: str
    mode: str = "General IEEE conference operations"
    knowledge_scope: str = "IEEE ITSS"
    conference_id: str | None = None
    conference_series: str | None = None


class SettingsUpdate(BaseModel):
    portfolio_start_year: int | None = None
    score_weights: dict[str, float] | None = None
    score_settings: dict[str, Any] | None = None
    status_mappings: dict[str, str] | None = None
    feature_flags: dict[str, bool] | None = None
    role_permissions: dict[str, dict[str, bool]] | None = None
    assistant_system_prompt: str | None = Field(default=None, min_length=80, max_length=5000)


class ReferenceConfigUpdate(BaseModel):
    reference_config: dict[str, Any]


class LlmTestMessage(BaseModel):
    message: str = Field(default="Reply with one short sentence confirming the IEEE ITSS dashboard LLM connection works.", min_length=1, max_length=500)


def assistant_system_prompt(session: Session) -> str:
    setting = session.get(DashboardSettings, "assistant_system_prompt")
    if not setting:
        return DEFAULT_ASSISTANT_SYSTEM_PROMPT
    try:
        value = json.loads(setting.value_json)
    except json.JSONDecodeError:
        return DEFAULT_ASSISTANT_SYSTEM_PROMPT
    return value if isinstance(value, str) and value.strip() else DEFAULT_ASSISTANT_SYSTEM_PROMPT


def finance_contact_name(conference: Conference) -> str | None:
    active_contacts = [contact for contact in conference.contacts if contact.active]
    for contact in active_contacts:
        if contact.role == "Finance Chair":
            return contact.name
    for contact in active_contacts:
        if "finance" in contact.role.lower():
            return contact.name
    return conference.financial_analyst


def conference_payload(conference: Conference) -> dict[str, Any]:
    details = json.loads(conference.score_details_json or "{}")
    source_details = json.loads(conference.source_details_json or "{}")
    return {
        "id": conference.id,
        "conference_number": normalize_record_number(conference.conference_number),
        "acronym": conference.acronym,
        "year": conference.year,
        "official_title": conference.official_title,
        "canonical_name": conference.canonical_name,
        "conference_series": conference.conference_series,
        "conference_category": conference.conference_category,
        "sponsorship_type": conference.sponsorship_type,
        "parent_conference_id": conference.parent_conference_id,
        "lifecycle_phase": conference.lifecycle_phase,
        "suggested_phase": conference.suggested_phase,
        "phase_override": conference.phase_override,
        "phase_differs": conference.lifecycle_phase != conference.suggested_phase,
        "conference_status": conference.conference_status,
        "active": conference.active,
        "start_date": conference.start_date.isoformat() if conference.start_date else None,
        "end_date": conference.end_date.isoformat() if conference.end_date else None,
        "city": conference.city,
        "country": conference.country,
        "website": conference.website,
        "estimated_attendees": conference.estimated_attendees,
        "actual_attendees": conference.actual_attendees,
        "last_source_update": conference.last_source_update.isoformat() if conference.last_source_update else None,
        "last_reviewed_date": conference.last_reviewed_date.isoformat() if conference.last_reviewed_date else None,
        "comments": conference.comments,
        "source_details": source_details,
        "application_status": conference.application_status,
        "application_status_raw": conference.application_status_raw,
        "application_submitted_date": conference.application_submitted_date.isoformat() if conference.application_submitted_date else None,
        "application_approved_date": conference.application_approved_date.isoformat() if conference.application_approved_date else None,
        "mou_status": conference.mou_status,
        "mou_status_raw": conference.mou_status_raw,
        "mou_signed_date": conference.mou_signed_date.isoformat() if conference.mou_signed_date else None,
        "finance_status": conference.finance_status,
        "financial_analyst": finance_contact_name(conference),
        "committee_contact": conference.committee_contact,
        "currency": conference.currency,
        "total_income_current": conference.total_income_current,
        "total_expense_current": conference.total_expense_current,
        "budgeted_income_total": conference.budgeted_income_total,
        "budgeted_expense_total": conference.budgeted_expense_total,
        "itss_loan_requested": conference.itss_loan_requested,
        "itss_loan_amount": conference.itss_loan_amount,
        "accounting_close_date": conference.accounting_close_date.isoformat() if conference.accounting_close_date else None,
        "publication_status": conference.publication_status,
        "proceedings_submitted_date": conference.proceedings_submitted_date.isoformat() if conference.proceedings_submitted_date else None,
        "xplore_posting_date": conference.xplore_posting_date.isoformat() if conference.xplore_posting_date else None,
        "score": conference.score,
        "base_score": conference.base_score,
        "issue_penalty": conference.issue_penalty,
        "data_completeness": conference.data_completeness,
        "status_band": conference.status_band,
        "score_details": details,
        "open_issue_count": sum(1 for issue in conference.issues if issue.active),
        "contacts": [
            {
                "id": c.id,
                "role": c.role,
                "name": c.name,
                "email": c.email,
                "organization": c.organization,
                "phone": c.phone,
                "is_primary": c.is_primary,
                "active": c.active,
            }
            for c in conference.contacts
        ],
        "milestones": [
            {
                "id": m.id,
                "code": m.definition.code,
                "name": m.definition.name,
                "dimension": m.definition.score_dimension,
                "status": m.status,
                "due_date": m.due_date.isoformat() if m.due_date else None,
                "completed_date": m.completed_date.isoformat() if m.completed_date else None,
                "manual_override": m.manual_override,
                "comments": m.comments,
                "last_updated": m.last_updated.isoformat(),
            }
            for m in conference.milestones
        ],
        "comments_history": [
            {
                "id": item.id,
                "comment": item.comment,
                "author": item.author,
                "created_at": item.created_at.isoformat(),
                "updated_at": item.updated_at.isoformat(),
            }
            for item in sorted(conference.comments_history, key=lambda row: row.updated_at, reverse=True)
        ],
    }


def issue_payload(issue: Issue, conference: Conference | None = None) -> dict[str, Any]:
    return {
        "id": issue.id,
        "conference_id": issue.conference_id,
        "conference_name": conference.canonical_name if conference else None,
        "issue_key": issue.issue_key,
        "title": issue.title,
        "description": issue.description,
        "category": issue.category,
        "severity": issue.severity,
        "issue_status": issue.issue_status,
        "review_assessment": issue.review_assessment,
        "source_type": issue.source_type,
        "owner": issue.owner,
        "date_detected": issue.date_detected.isoformat(),
        "due_date": issue.due_date.isoformat() if issue.due_date else None,
        "ai_recommendation": issue.ai_recommendation,
        "active": issue.active,
    }


router = APIRouter()


def enforce_reference_config(session: Session, config: dict[str, Any]) -> dict[str, int]:
    changed = {
        "conferences": 0,
        "milestones": 0,
        "contacts": 0,
        "issues": 0,
        "mappings": 0,
    }

    def allowed(key: str) -> set[str]:
        return set(config.get(key, []))

    phase_allowed = allowed("lifecycle_phases")
    conference_status_allowed = allowed("conference_statuses")
    status_allowed = allowed("normalized_statuses")
    sponsorship_allowed = allowed("sponsorship_types")
    contact_role_allowed = allowed("contact_roles")
    issue_category_allowed = allowed("issue_categories")
    issue_severity_allowed = allowed("issue_severities")
    review_allowed = allowed("review_assessments")
    review_fallback = "Unreviewed" if "Unreviewed" in review_allowed else next(iter(review_allowed), "Unreviewed")
    series_allowed = {
        value
        for item in config.get("conference_series", [])
        if isinstance(item, dict)
        for value in (str(item.get("code", "")), str(item.get("name", "")))
        if value
    }

    for conference in session.scalars(select(Conference)):
        touched = False
        for field, allowed_values in {
            "lifecycle_phase": phase_allowed,
            "suggested_phase": phase_allowed,
            "sponsorship_type": sponsorship_allowed,
            "conference_series": series_allowed,
            "conference_status": conference_status_allowed,
            "application_status": status_allowed,
            "mou_status": status_allowed,
            "finance_status": status_allowed,
            "publication_status": status_allowed,
        }.items():
            if getattr(conference, field) not in allowed_values:
                setattr(conference, field, "Unknown")
                touched = True
        if touched:
            changed["conferences"] += 1
            sync_conference_facts_from_milestones(conference)
            recalculate(conference, session)

    for milestone in session.scalars(select(ConferenceMilestone)):
        if milestone.status not in status_allowed:
            milestone.status = "Unknown"
            milestone.last_updated = now()
            changed["milestones"] += 1

    for contact in session.scalars(select(Contact)):
        if contact.role not in contact_role_allowed:
            contact.role = "Unknown"
            changed["contacts"] += 1

    for issue in session.scalars(select(Issue)):
        touched = False
        if issue.category not in issue_category_allowed:
            issue.category = "Unknown"
            touched = True
        if issue.severity not in issue_severity_allowed:
            issue.severity = "Unknown"
            touched = True
        if issue.review_assessment not in review_allowed:
            issue.review_assessment = review_fallback
            touched = True
        if issue.issue_status not in status_allowed:
            issue.issue_status = "Unknown"
            touched = True
        if touched:
            changed["issues"] += 1

    for mapping in session.scalars(select(StatusMapping)):
        if mapping.normalized_value not in status_allowed:
            mapping.normalized_value = "Unknown"
            changed["mappings"] += 1

    session.flush()
    for conference in session.scalars(select(Conference)):
        sync_conference_facts_from_milestones(conference)
        detect_issues(conference, session)
        recalculate(conference, session)
    return changed


def configured_value(session: Session, key: str, value: str | None) -> str:
    if value is None:
        if key == "review_assessments":
            return "Unreviewed"
        return "UNKNOWN" if key == "conference_series" else "Unknown"
    cleaned = " ".join(str(value).strip().split())
    if key == "conference_series":
        allowed = allowed_reference_values(session, key)
        if cleaned in allowed:
            return cleaned
        upper = cleaned.upper()
        return upper if upper in allowed else "UNKNOWN"
    if cleaned in allowed_reference_values(session, key):
        return cleaned
    if key == "review_assessments":
        return "Unreviewed"
    return "Unknown"


@router.get("/api/dashboard/summary")
def dashboard_summary(session: Session = Depends(get_session)) -> dict[str, Any]:
    conferences = list(session.scalars(select(Conference).where(Conference.active.is_(True))))
    issues = list(session.scalars(select(Issue).where(Issue.active.is_(True))))
    flagship_cards: list[dict[str, Any]] = []
    flagship_groups: dict[str, list[dict[str, Any]]] = {"ITSC": [], "IV": []}
    for conf in sorted(conferences, key=lambda item: (item.normalized_acronym, item.year), reverse=True):
        if conf.normalized_acronym not in flagship_groups:
            continue
        item = conference_payload(conf)
        flagship_cards.append(item)
        flagship_groups[conf.normalized_acronym].append(item)
    surplus_pcts = []
    for conf in conferences:
        income = conf.total_income_current
        expense = conf.total_expense_current
        if income is None or expense is None:
            income = conf.budgeted_income_total
            expense = conf.budgeted_expense_total
        if income is not None and expense is not None and expense != 0:
            surplus_pct = ((income - expense) / expense) * 100
            if abs(surplus_pct) > 0.01:
                surplus_pcts.append(surplus_pct)
    return {
        "conference_count": len(conferences),
        "open_issue_count": len(issues),
        "critical_issue_count": sum(1 for item in issues if item.severity == "Critical"),
        "average_score": round(sum(c.score for c in conferences) / len(conferences), 1) if conferences else 0,
        "average_surplus_percentage": round(sum(surplus_pcts) / len(surplus_pcts), 1) if surplus_pcts else None,
        "status_counts": dict(Counter(c.conference_status for c in conferences)),
        "health_counts": dict(Counter(c.status_band for c in conferences)),
        "phase_counts": dict(Counter(c.lifecycle_phase for c in conferences)),
        "flagship_cards": flagship_cards,
        "flagship_groups": flagship_groups,
        "last_source_update": max((c.last_source_update for c in conferences if c.last_source_update), default=None),
        "azure_openai": azure_status(mask=True),
        "embeddings": embedding_status(mask=True),
        "onboarding": len(conferences) == 0,
    }


@router.get("/api/reference-data")
def reference_data(session: Session = Depends(get_session)) -> dict[str, Any]:
    config = reference_config(session)
    mappings = {m.source_value: m.normalized_value for m in session.scalars(select(StatusMapping).where(StatusMapping.active.is_(True)))}
    return {
        "committee_members": config["committee_members"],
        "ieee_brand": {"blue": IEEE_BLUE, "teal": IEEE_TEAL, "amber": IEEE_AMBER, "orange": IEEE_ORANGE, "red": IEEE_RED, "gray": IEEE_GRAY},
        "conference_series": config["conference_series"],
        "sponsorship_types": config["sponsorship_types"],
        "lifecycle_phases": config["lifecycle_phases"],
        "normalized_statuses": config["normalized_statuses"],
        "contact_roles": config["contact_roles"],
        "issue_categories": config["issue_categories"],
        "issue_severities": config["issue_severities"],
        "review_assessments": config["review_assessments"],
        "status_mappings": mappings,
    }


@router.get("/api/conferences")
def list_conferences(
    q: str | None = None,
    year: int | None = None,
    series: str | None = None,
    include_archived: bool = False,
    session: Session = Depends(get_session),
) -> dict[str, Any]:
    query = select(Conference)
    if not include_archived:
        query = query.where(Conference.active.is_(True))
    if year:
        query = query.where(Conference.year == year)
    if series:
        query = query.where(Conference.conference_series == series)
    items = list(session.scalars(query.order_by(Conference.year.desc(), Conference.acronym.asc())))
    if q:
        needle = q.lower()
        alias_matches = {
            alias.conference_id
            for alias in session.scalars(select(Alias).where(Alias.alias.ilike(f"%{q}%"), Alias.active.is_(True)))
        }
        items = [
            c
            for c in items
            if needle in c.canonical_name.lower()
            or needle in c.official_title.lower()
            or (c.conference_number and needle in c.conference_number.lower())
            or (c.city and needle in c.city.lower())
            or (c.country and needle in c.country.lower())
            or c.id in alias_matches
        ]
    return {"items": [conference_payload(c) for c in items]}


@router.post("/api/conferences", status_code=201)
def create_conference(payload: ConferenceIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    norm = normalize_acronym(payload.acronym)
    if payload.conference_number:
        duplicate_number = session.scalar(select(Conference).where(Conference.conference_number == payload.conference_number))
        if duplicate_number:
            raise HTTPException(409, "A conference with this IEEE Conference Record Number already exists.")
    duplicate = session.scalar(
        select(Conference).where(
            Conference.normalized_acronym == norm,
            Conference.year == payload.year,
            Conference.parent_conference_id.is_(None) if payload.parent_conference_id is None else Conference.parent_conference_id == payload.parent_conference_id,
        )
    )
    if duplicate:
        raise HTTPException(409, "A conference with this acronym and year already exists. Use a child event to track sub-events.")
    conference = Conference(
        conference_number=payload.conference_number,
        acronym=payload.acronym.upper(),
        normalized_acronym=norm,
        year=payload.year,
        official_title=payload.official_title,
        canonical_name=f"{payload.acronym.upper()} {payload.year}",
        conference_series=configured_value(session, "conference_series", payload.conference_series),
        conference_category=payload.conference_category,
        sponsorship_type=configured_value(session, "sponsorship_types", payload.sponsorship_type),
        parent_conference_id=payload.parent_conference_id,
        lifecycle_phase=configured_value(session, "lifecycle_phases", payload.lifecycle_phase),
        phase_override=False,
        conference_status="Unknown",
        start_date=payload.start_date,
        end_date=payload.end_date,
        submission_deadline=payload.submission_deadline,
        notification_date=payload.notification_date,
        camera_ready_deadline=payload.camera_ready_deadline,
        city=payload.city,
        state_province=payload.state_province,
        country=payload.country,
        ieee_region=payload.ieee_region,
        venue=payload.venue,
        website=payload.website,
        estimated_attendees=payload.estimated_attendees,
    )
    session.add(conference)
    session.flush()
    if payload.primary_contact or payload.primary_contact_email:
        session.add(
            Contact(
                conference_id=conference.id,
                role="Information Contact",
                name=payload.primary_contact or "Primary contact",
                email=str(payload.primary_contact_email) if payload.primary_contact_email else None,
                is_primary=True,
            )
        )
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="create")
    session.commit()
    return conference_payload(conference)


@router.get("/api/conferences/{conference_id}")
def get_conference(conference_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    return conference_payload(require_conference(session, conference_id))


@router.patch("/api/conferences/{conference_id}")
def update_conference(conference_id: str, payload: ConferenceUpdate, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    data = payload.model_dump(exclude_unset=True)
    comment = data.pop("change_comment", None)
    for derived_field in (
        "conference_status",
        "lifecycle_phase",
        "phase_override",
        "application_status",
        "mou_status",
        "finance_status",
        "publication_status",
    ):
        data.pop(derived_field, None)
    for field, value in data.items():
        old = getattr(conference, field)
        if old != value:
            setattr(conference, field, value)
            session.add(
                FieldChange(
                    conference_id=conference.id,
                    field_name=field,
                    old_value=str(old) if old is not None else None,
                    new_value=str(value) if value is not None else None,
                    change_type="Manual",
                    source="UI",
                    comment=comment,
                )
            )
    if "conference_number" in data and data["conference_number"]:
        duplicate = session.scalar(select(Conference).where(Conference.conference_number == data["conference_number"], Conference.id != conference.id))
        if duplicate:
            raise HTTPException(409, "Another conference already uses this IEEE Conference Record Number.")
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="manual update")
    session.commit()
    return conference_payload(conference)


@router.post("/api/conferences/{conference_id}/refresh")
def refresh_conference(conference_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    sync_conference_facts_from_milestones(conference)
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="manual refresh")
    session.commit()
    return conference_payload(conference)


@router.patch("/api/conferences/{conference_id}/milestones/{milestone_id}")
def update_conference_milestone(conference_id: str, milestone_id: str, payload: MilestoneUpdate, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    milestone = session.get(ConferenceMilestone, milestone_id)
    if not milestone or milestone.conference_id != conference.id:
        raise HTTPException(404, "Unknown conference milestone.")
    milestone.status = configured_value(session, "normalized_statuses", payload.status)
    if "due_date" in payload.model_fields_set:
        milestone.due_date = payload.due_date
    milestone.comments = payload.comments
    milestone.manual_override = True
    milestone.last_updated = now()
    sync_conference_facts_from_milestones(conference)
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="milestone status update")
    session.commit()
    return conference_payload(conference)


@router.post("/api/conferences/{conference_id}/contacts", status_code=201)
def create_contact(conference_id: str, payload: ContactIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    if payload.is_primary:
        for contact in conference.contacts:
            contact.is_primary = False
    contact = Contact(
        conference_id=conference.id,
        role=configured_value(session, "contact_roles", payload.role),
        name=payload.name,
        email=str(payload.email) if payload.email else None,
        organization=payload.organization,
        phone=payload.phone,
        is_primary=payload.is_primary,
    )
    session.add(contact)
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="contact create")
    session.commit()
    return conference_payload(conference)


@router.patch("/api/conferences/{conference_id}/contacts/{contact_id}")
def update_contact(conference_id: str, contact_id: str, payload: ContactUpdate, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    contact = session.get(Contact, contact_id)
    if not contact or contact.conference_id != conference.id:
        raise HTTPException(404, "Unknown conference contact.")
    data = payload.model_dump(exclude_unset=True)
    if data.get("is_primary"):
        for item in conference.contacts:
            item.is_primary = False
    for field, value in data.items():
        if field == "role":
            value = configured_value(session, "contact_roles", value)
        setattr(contact, field, str(value) if field == "email" and value else value)
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="contact update")
    session.commit()
    return conference_payload(conference)


@router.delete("/api/conferences/{conference_id}/contacts/{contact_id}")
def delete_contact(conference_id: str, contact_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    contact = session.get(Contact, contact_id)
    if not contact or contact.conference_id != conference.id:
        raise HTTPException(404, "Unknown conference contact.")
    session.delete(contact)
    detect_issues(conference, session)
    recalculate(conference, session)
    create_snapshot(conference, session, reason="contact delete")
    session.commit()
    return conference_payload(conference)


@router.post("/api/conferences/{conference_id}/comments", status_code=201)
def add_conference_comment(conference_id: str, payload: ConferenceCommentIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    conference.comments = payload.comment
    conference.last_reviewed_date = date.today()
    row = ConferenceComment(conference_id=conference.id, comment=payload.comment)
    session.add(row)
    session.commit()
    return conference_payload(conference)


@router.patch("/api/conferences/{conference_id}/comments/{comment_id}")
def update_conference_comment(conference_id: str, comment_id: str, payload: ConferenceCommentIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    row = session.get(ConferenceComment, comment_id)
    if not row or row.conference_id != conference.id:
        raise HTTPException(404, "Unknown conference comment.")
    row.comment = payload.comment
    conference.comments = payload.comment
    conference.last_reviewed_date = date.today()
    session.commit()
    return conference_payload(conference)


@router.delete("/api/conferences/{conference_id}/comments/{comment_id}")
def delete_conference_comment(conference_id: str, comment_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    row = session.get(ConferenceComment, comment_id)
    if not row or row.conference_id != conference.id:
        raise HTTPException(404, "Unknown conference comment.")
    session.delete(row)
    session.flush()
    latest = session.scalar(
        select(ConferenceComment)
        .where(ConferenceComment.conference_id == conference.id)
        .order_by(ConferenceComment.updated_at.desc())
    )
    conference.comments = latest.comment if latest else None
    conference.last_reviewed_date = latest.updated_at.date() if latest else None
    session.commit()
    return conference_payload(conference)


@router.post("/api/conferences/{conference_id}/archive")
def archive_conference(conference_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    conference = require_conference(session, conference_id)
    conference.active = False
    session.commit()
    return {"status": "archived"}


@router.post("/api/conferences/{conference_id}/restore")
def restore_conference(conference_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    conference = require_conference(session, conference_id)
    conference.active = True
    session.commit()
    return {"status": "restored"}


@router.delete("/api/conferences/{conference_id}")
def delete_conference_record(conference_id: str, payload: ConferenceDeleteIn, session: Session = Depends(get_session)) -> dict[str, str]:
    conference = require_conference(session, conference_id)
    expected = normalize_record_number(conference.conference_number) or "NO RECORD"
    provided = normalize_record_number(payload.confirmation_record_number) or payload.confirmation_record_number.strip()
    if provided != expected:
        raise HTTPException(400, "Conference record number confirmation did not match.")
    issue_ids = [issue.id for issue in session.scalars(select(Issue).where(Issue.conference_id == conference.id))]
    if issue_ids:
        session.execute(delete(IssueComment).where(IssueComment.issue_id.in_(issue_ids)))
    for model in (Alias, ConferenceMilestone, Issue, ConferenceComment, FieldChange, Snapshot, ScoreHistory, EmailDraft, DashboardPin):
        session.execute(delete(model).where(model.conference_id == conference.id))
    session.execute(delete(Document).where(Document.conference_id == conference.id))
    for child in session.scalars(select(Conference).where(Conference.parent_conference_id == conference.id)):
        child.parent_conference_id = None
    session.delete(conference)
    session.commit()
    return {"status": "deleted"}


@router.get("/api/search")
def global_search(q: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    q = q.strip()
    conferences = list_conferences(q=q, include_archived=True, session=session)["items"]
    issues = [
        issue_payload(issue, session.get(Conference, issue.conference_id))
        for issue in session.scalars(select(Issue).where(Issue.title.ilike(f"%{q}%") | Issue.description.ilike(f"%{q}%")))
    ]
    contacts = [
        {"id": c.id, "conference_id": c.conference_id, "name": c.name, "email": c.email, "role": c.role}
        for c in session.scalars(select(Contact).where(Contact.name.ilike(f"%{q}%") | Contact.email.ilike(f"%{q}%")))
    ]
    documents = [
        {"id": d.id, "title": d.title, "category": d.document_category, "scope": d.knowledge_scope}
        for d in session.scalars(select(Document).where(Document.title.ilike(f"%{q}%") | Document.extracted_text.ilike(f"%{q}%")))
    ]
    return {"conferences": conferences, "issues": issues, "contacts": contacts, "documents": documents}


@router.get("/api/conferences/{conference_id}/history")
def conference_history(conference_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    require_conference(session, conference_id)
    changes = list(session.scalars(select(FieldChange).where(FieldChange.conference_id == conference_id).order_by(FieldChange.timestamp.desc())))
    snapshots = list(session.scalars(select(Snapshot).where(Snapshot.conference_id == conference_id).order_by(Snapshot.snapshot_timestamp.desc())))
    scores = list(session.scalars(select(ScoreHistory).where(ScoreHistory.conference_id == conference_id).order_by(ScoreHistory.created_at.desc()).limit(24)))
    return {
        "changes": [
            {"field": c.field_name, "old": c.old_value, "new": c.new_value, "source": c.source, "timestamp": c.timestamp.isoformat(), "comment": c.comment}
            for c in changes
        ],
        "snapshots": [{"id": s.id, "timestamp": s.snapshot_timestamp.isoformat(), "payload": json.loads(s.payload_json)} for s in snapshots],
        "score_history": [{"score": s.score, "data_completeness": s.data_completeness, "created_at": s.created_at.isoformat()} for s in scores],
    }


@router.post("/api/conferences/{conference_id}/snapshot")
def snapshot_endpoint(conference_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    snapshot = create_snapshot(require_conference(session, conference_id), session, reason="explicit")
    session.commit()
    return {"snapshot_id": snapshot.id}


@router.get("/api/issues")
def list_issues(conference_id: str | None = None, assessment: str | None = None, session: Session = Depends(get_session)) -> dict[str, Any]:
    query = select(Issue)
    if conference_id:
        query = query.where(Issue.conference_id == conference_id)
    if assessment:
        query = query.where(Issue.review_assessment == assessment)
    issues = list(session.scalars(query.order_by(Issue.date_detected.desc())))
    return {"items": [issue_payload(issue, session.get(Conference, issue.conference_id)) for issue in issues]}


@router.post("/api/issues", status_code=201)
def create_issue(payload: IssueIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, payload.conference_id)
    issue = Issue(
        conference_id=conference.id,
        issue_key=f"{conference.canonical_name}-MANUAL-{uuid.uuid4().hex[:6].upper()}",
        title=payload.title,
        description=payload.description,
        category=configured_value(session, "issue_categories", payload.category),
        severity=configured_value(session, "issue_severities", payload.severity),
        owner=payload.owner,
        due_date=payload.due_date,
        source_type="Manual",
    )
    session.add(issue)
    recalculate(conference, session)
    session.commit()
    return issue_payload(issue, conference)


@router.patch("/api/issues/{issue_id}")
def update_issue(issue_id: str, payload: IssueUpdate, session: Session = Depends(get_session)) -> dict[str, Any]:
    issue = require_issue(session, issue_id)
    data = payload.model_dump(exclude_unset=True)
    reference_fields = {
        "issue_status": "normalized_statuses",
        "review_assessment": "review_assessments",
        "category": "issue_categories",
        "severity": "issue_severities",
    }
    data.pop("user_comment", None)
    for field, value in data.items():
        if field in reference_fields:
            value = configured_value(session, reference_fields[field], value)
        setattr(issue, field, value)
    if issue.issue_status == "Resolved" or issue.review_assessment == "Not an Issue":
        issue.resolution_date = date.today()
        issue.active = False
    conference = require_conference(session, issue.conference_id)
    recalculate(conference, session)
    session.commit()
    return issue_payload(issue, conference)


@router.post("/api/issues/{issue_id}/resolve")
def resolve_issue(issue_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    issue = require_issue(session, issue_id)
    issue.issue_status = "Resolved"
    issue.review_assessment = "Not an Issue"
    issue.resolution_date = date.today()
    issue.active = False
    conference = require_conference(session, issue.conference_id)
    recalculate(conference, session)
    session.commit()
    return issue_payload(issue, conference)


@router.delete("/api/issues/{issue_id}")
def delete_issue(issue_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    issue = require_issue(session, issue_id)
    conference = require_conference(session, issue.conference_id)
    session.delete(issue)
    recalculate(conference, session)
    session.commit()
    return {"status": "deleted"}


@router.post("/api/issues/{issue_id}/comments")
def add_issue_comment(issue_id: str, comment: str = Form(...), session: Session = Depends(get_session)) -> dict[str, str]:
    require_issue(session, issue_id)
    row = IssueComment(issue_id=issue_id, comment=comment)
    session.add(row)
    session.commit()
    return {"comment_id": row.id}


@router.post("/api/issues/{issue_id}/ai-recommendation")
def issue_recommendation(issue_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    issue = require_issue(session, issue_id)
    conference = require_conference(session, issue.conference_id)
    prompt = (
        f"Recommend next actions for IEEE ITSS conference issue. Conference: {conference.canonical_name}; "
        f"phase: {conference.lifecycle_phase}; issue: {issue.title}; severity: {issue.severity}; details: {issue.description}. "
        "Use only the supplied facts and do not invent policy."
    )
    recommendation = call_azure_chat(prompt) or local_recommendation(conference, issue)
    issue.ai_recommendation = recommendation
    issue.recommendation_generated_date = now()
    session.commit()
    return issue_payload(issue, conference)


@router.post("/api/conferences/{conference_id}/generate-issues")
def generate_conference_issues(conference_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, conference_id)
    payload = conference_payload(conference)
    sources = retrieve_sources(
        ChatIn(
            question="IEEE conference governance finance publication MOU operations registration closure issue detection guidance",
            knowledge_scope="IEEE ITSS",
            conference_id=conference.id,
            conference_series=conference.conference_series,
        ),
        session,
    )
    prompt = json.dumps(
        {
            "conference": {
                key: payload.get(key)
                for key in (
                    "conference_number",
                    "canonical_name",
                    "official_title",
                    "lifecycle_phase",
                    "conference_status",
                    "start_date",
                    "end_date",
                    "city",
                    "country",
                    "mou_status",
                    "application_status",
                    "finance_status",
                    "publication_status",
                    "accounting_close_date",
                    "xplore_posting_date",
                    "estimated_attendees",
                    "actual_attendees",
                    "comments",
                )
            },
            "milestones": payload.get("milestones", []),
            "recent_comments": payload.get("comments_history", [])[:5],
            "knowledge_sources": sources,
            "instructions": (
                "Identify actionable open issues for IEEE ITSS conference management. Use conference facts, milestones, comments, and knowledge excerpts. "
                "Do not create issues for items already complete or not applicable. Return strict JSON with key 'issues'. "
                "Each issue must include title, description, category, severity, review_assessment, owner, due_date. "
                "Allowed severities: Informational, Low, Medium, High, Critical. "
                "Allowed review_assessment values: Unreviewed, On Track, Needs Follow-up, Not an Issue."
            ),
        },
        default=str,
    )
    try:
        raw = llm_chat_completion_text(
            system_prompt="You are an IEEE ITSS conference operations reviewer. Return only valid JSON.",
            user_prompt=prompt,
            temperature=0.1,
            max_tokens=2000,
            response_format={"type": "json_object"},
        )
    except RuntimeError as exc:
        raise HTTPException(503, f"LLM issue generation failed: {exc}") from exc
    generated = parse_generated_issues(raw, session)
    created: list[Issue] = []
    for item in generated:
        issue = Issue(
            conference_id=conference.id,
            issue_key=f"{conference.canonical_name}-LLM-{uuid.uuid4().hex[:8]}",
            title=item.title,
            description=item.description,
            category=configured_value(session, "issue_categories", item.category),
            severity=configured_value(session, "issue_severities", item.severity),
            issue_status="Open",
            review_assessment=configured_value(session, "review_assessments", item.review_assessment),
            source_type="LLM",
            source_field="conference_review",
            rule_identifier="LLM_GENERATED",
            owner=item.owner,
            due_date=item.due_date,
        )
        session.add(issue)
        created.append(issue)
    recalculate(conference, session)
    session.commit()
    return {"created": len(created), "items": [issue_payload(issue, conference) for issue in created]}


@router.get("/api/imports/template.xlsx")
def excel_template() -> Response:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame([{"report_date": date.today().isoformat(), "report_title": "IEEE ITSS Portfolio Status", "source": "Local", "prepared_by": "", "notes": ""}]).to_excel(writer, index=False, sheet_name="Metadata")
        pd.DataFrame(columns=CONFERENCE_COLUMNS).to_excel(writer, index=False, sheet_name="Conferences")
        pd.DataFrame(columns=CONTACT_COLUMNS).to_excel(writer, index=False, sheet_name="Contacts")
        pd.DataFrame(columns=ISSUE_COLUMNS).to_excel(writer, index=False, sheet_name="Issues")
    return Response(output.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=itss-status-template.xlsx"})


@router.get("/api/imports/template.csv")
def csv_template() -> Response:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(CONFERENCE_COLUMNS)
    return Response(output.getvalue(), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=itss-conference-status-template.csv"})


@router.get("/api/imports/field-guide")
def field_guide() -> Response:
    lines = [
        "IEEE ITSS import field guide",
        "",
        "Use the CSV template from /api/imports/template.csv. The first row must contain the canonical column names.",
        "Required for every row: acronym, year, official_title.",
        "Required for new conferences: acronym, year, official_title, conference_series, sponsorship_type, lifecycle_phase.",
        "Matching: conference_number is used first when present; otherwise acronym + year is used.",
        "Blank cells leave existing values unchanged. [CLEAR] clears a value.",
        "Dates must use ISO format: YYYY-MM-DD, for example 2027-06-22. Values like 22-Jun-2027 are rejected.",
        "Numbers must be plain numbers. Currency fields may include comma thousands separators, for example 12,500.00.",
        "Use configured status/reference values where possible, such as Approved, Complete, Submitted, In Progress, Blocked, Unknown.",
        "",
        "Conference columns:",
    ]
    lines.extend(f"- {column}" for column in CONFERENCE_COLUMNS)
    return Response("\n".join(lines), media_type="text/plain")


@router.post("/api/imports/validate")
async def validate_import(file: UploadFile = File(...), session: Session = Depends(get_session)) -> dict[str, Any]:
    data = await file.read()
    preview = build_import_preview(file.filename or "upload", data, session)
    return preview


@router.post("/api/imports/apply")
async def apply_import(file: UploadFile = File(...), selected_changes_json: str | None = Form(default=None), session: Session = Depends(get_session)) -> dict[str, Any]:
    data = await file.read()
    preview = build_import_preview(file.filename or "upload", data, session)
    batch = ImportBatch(
        original_filename=file.filename or "upload",
        file_type=Path(file.filename or "").suffix.lower().lstrip("."),
        file_hash=hashlib.sha256(data).hexdigest(),
        rows_count=preview["summary"]["rows"],
        new_count=preview["summary"]["new"],
        changed_count=preview["summary"]["changed"],
        unchanged_count=preview["summary"]["unchanged"],
        conflict_count=preview["summary"]["conflicts"],
        import_status="Applied",
        preview_json=json.dumps(jsonable_encoder(preview)),
        file_data=data,
    )
    session.add(batch)
    selected = json.loads(selected_changes_json) if selected_changes_json else None
    applied = 0
    skipped = 0
    for row in preview["rows"]:
        selected_fields = selected_import_fields(selected, row)
        if selected is not None and selected_fields is None:
            continue
        if selected is not None and not selected_fields:
            continue
        if row["validation_result"] != "valid":
            if selected is None or not can_apply_partial_import_row(row, selected_fields or set()):
                skipped += 1
                continue
        conference = apply_import_row(row["source"], session, batch, selected_fields=selected_fields)
        detect_issues(conference, session)
        recalculate(conference, session)
        create_snapshot(conference, session, reason="import")
        applied += 1
    milestone_applied = 0
    milestone_approved = selected is None or selected.get("__milestones__") is not None
    for mrow in preview.get("milestone_rows", []):
        if not milestone_approved:
            continue
        if mrow["errors"]:
            skipped += 1
            continue
        if not mrow["changes"]:
            continue
        m_conference = match_conference(session, mrow["conference_number"] or None, mrow["acronym"], mrow["year"])
        if not m_conference:
            continue
        m_definition = session.scalar(select(MilestoneDefinition).where(MilestoneDefinition.code == mrow["milestone_code"]))
        if not m_definition:
            continue
        m_milestone = session.scalar(
            select(ConferenceMilestone).where(
                ConferenceMilestone.conference_id == m_conference.id,
                ConferenceMilestone.definition_id == m_definition.id,
            )
        )
        if not m_milestone:
            continue
        for mchange in mrow["changes"]:
            mfield = mchange["field"]
            mnew = mchange["new"]
            if comparable_import_value(getattr(m_milestone, mfield)) != comparable_import_value(mnew):
                setattr(m_milestone, mfield, mnew)
                m_milestone.manual_override = True if mfield == "status" else m_milestone.manual_override
                m_milestone.last_updated = now()
                session.add(FieldChange(conference_id=m_conference.id, entity="ConferenceMilestone", field_name=f"milestone_{mfield}", old_value=str(mchange.get("old", "")) if mchange.get("old") else None, new_value=str(mnew) if mnew else None, change_type="Import", source=batch.original_filename, import_batch_id=batch.id))
        detect_issues(m_conference, session)
        recalculate(m_conference, session)
        create_snapshot(m_conference, session, reason="milestone import")
        milestone_applied += 1
    if applied == 0 and milestone_applied == 0:
        session.rollback()
        raise HTTPException(409, "Select at least one valid new or changed row before applying changes.")
    session.commit()
    return {"batch_id": batch.id, "applied_rows": applied, "milestone_applied": milestone_applied, "skipped_rows": skipped, "summary": preview["summary"]}


@router.post("/api/imports/{batch_id}/rollback")
def rollback_import(batch_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    batch = session.get(ImportBatch, batch_id)
    if not batch:
        raise HTTPException(404, "Unknown import batch.")
    changes = list(session.scalars(select(FieldChange).where(FieldChange.import_batch_id == batch_id).order_by(FieldChange.timestamp.desc())))
    restored = 0
    for change in changes:
        conference = session.get(Conference, change.conference_id)
        if conference and hasattr(conference, change.field_name):
            setattr(conference, change.field_name, change.old_value)
            restored += 1
    batch.import_status = "Rolled Back"
    session.commit()
    return {"status": "rolled_back", "restored_fields": restored}


@router.get("/api/exports/portfolio.xlsx")
def export_portfolio(session: Session = Depends(get_session)) -> Response:
    output = io.BytesIO()
    conference_rows = [conference_export_row(c) for c in session.scalars(select(Conference).order_by(Conference.year.desc()))]
    issues = [issue_payload(i, session.get(Conference, i.conference_id)) for i in session.scalars(select(Issue))]
    contacts = [
        {
            "conference_number": c.conference_number,
            "acronym": c.acronym,
            "year": c.year,
            "contact_role": contact.role,
            "name": contact.name,
            "email": contact.email,
            "organization": contact.organization,
            "phone": contact.phone,
            "is_primary": contact.is_primary,
        }
        for c in session.scalars(select(Conference).order_by(Conference.year.desc()))
        for contact in c.contacts
    ]
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame(conference_rows, columns=CONFERENCE_COLUMNS).to_excel(writer, index=False, sheet_name="Conferences")
        pd.DataFrame(issues).to_excel(writer, index=False, sheet_name="Issues")
        pd.DataFrame(contacts, columns=CONTACT_COLUMNS).to_excel(writer, index=False, sheet_name="Contacts")
        milestones = [
            {
                "conference_number": normalize_record_number(c.conference_number),
                "acronym": c.acronym,
                "year": c.year,
                "milestone_code": cm.definition.code,
                "milestone_name": cm.definition.name,
                "status": cm.status,
                "due_date": cm.due_date.isoformat() if cm.due_date else None,
                "comments": cm.comments,
            }
            for c in session.scalars(select(Conference).order_by(Conference.year.desc()))
            for cm in c.milestones
        ]
        pd.DataFrame(milestones, columns=MILESTONE_COLUMNS).to_excel(writer, index=False, sheet_name="Milestones")
        pd.DataFrame([{"dimension": k, "weight": v} for k, v in SCORE_WEIGHTS.items()]).to_excel(writer, index=False, sheet_name="Score Settings")
    return Response(output.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=itss-portfolio.xlsx"})


@router.get("/api/exports/executive.pdf")
def export_executive_pdf(session: Session = Depends(get_session)) -> Response:
    conferences = list(session.scalars(select(Conference).where(Conference.active.is_(True)).order_by(Conference.year.desc())))
    lines = ["IEEE ITSS Conference Portfolio Status", f"Generated: {date.today().isoformat()}", ""]
    for c in conferences:
        open_issues = sum(1 for i in c.issues if i.active)
        lines.append(f"{c.canonical_name} | {c.lifecycle_phase} | Score {c.score:.1f} | Coverage {c.data_completeness:.1f}% | Open issues {open_issues}")
    return Response(simple_pdf(lines), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=itss-executive-status.pdf"})


@router.post("/api/templates", status_code=201)
async def upload_template(
    file: UploadFile = File(...),
    template_name: str = Form(...),
    short_description: str = Form(""),
    category: str = Form("Unknown"),
    session: Session = Depends(get_session),
) -> dict[str, Any]:
    data = await file.read()
    if not data:
        raise HTTPException(422, "The selected template file is empty.")
    clean_name = template_name.strip()
    if not clean_name:
        raise HTTPException(422, "Template name is required.")
    original = file.filename or "template"
    stored_name = f"{uuid.uuid4().hex}-{safe_filename(original)}"
    template = TemplateFile(
        template_name=clean_name,
        short_description=short_description.strip(),
        category=category.strip() or "Unknown",
        template_type=template_type_from_filename(original),
        file_name=stored_name,
        original_filename=original,
        file_data=data,
    )
    session.add(template)
    session.commit()
    return template_payload(template)


@router.get("/api/templates")
def list_templates(session: Session = Depends(get_session)) -> dict[str, Any]:
    return {"items": [template_payload(item) for item in session.scalars(select(TemplateFile).order_by(TemplateFile.updated_at.desc()))]}


@router.get("/api/templates/{template_id}/download")
def download_template(template_id: str, session: Session = Depends(get_session)) -> Response:
    template = session.get(TemplateFile, template_id)
    if not template:
        raise HTTPException(404, "Unknown template.")
    if template.file_data is not None:
        return Response(
            content=template.file_data,
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe_filename(template.original_filename)}"'},
        )
    path = app_path("APP_TEMPLATE_PATH", "./data/templates") / template.file_name
    if not path.exists():
        raise HTTPException(404, "Stored template file is missing.")
    return FileResponse(path, filename=template.original_filename, media_type="application/octet-stream")


@router.delete("/api/templates/{template_id}")
def delete_template(template_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    template = session.get(TemplateFile, template_id)
    if not template:
        raise HTTPException(404, "Unknown template.")
    try:
        app_path("APP_TEMPLATE_PATH", "./data/templates").joinpath(template.file_name).unlink(missing_ok=True)
    except OSError:
        pass
    session.delete(template)
    session.commit()
    return {"status": "deleted"}


@router.post("/api/documents", status_code=201)
async def upload_document(
    file: UploadFile = File(...),
    title: str = Form(...),
    document_category: str = Form("Other"),
    knowledge_scope: str = Form("IEEE ITSS"),
    conference_id: str | None = Form(default=None),
    conference_series: str | None = Form(default=None),
    version: str | None = Form(default=None),
    source_url: str | None = Form(default=None),
    session: Session = Depends(get_session),
) -> dict[str, Any]:
    data = await file.read()
    document_dir = app_path("APP_DOCUMENT_PATH", "./data/documents")
    stored_name = f"{uuid.uuid4().hex}-{safe_filename(file.filename or 'document')}"
    (document_dir / stored_name).write_bytes(data)
    text = extract_text(file.filename or "", data)
    chunks = chunk_text(text)
    document_id = str(uuid.uuid4())
    try:
        vector_result = write_vector_chunks(document_id, chunks, require_embeddings=bool(chunks))
    except RuntimeError as exc:
        (document_dir / stored_name).unlink(missing_ok=True)
        raise HTTPException(502, f"Document was extracted, but embedding indexing failed: {exc}") from exc
    doc = Document(
        id=document_id,
        title=title,
        file_name=stored_name,
        document_category=document_category,
        knowledge_scope=knowledge_scope,
        conference_id=conference_id,
        conference_series=conference_series,
        version=version,
        source_url=source_url,
        page_count=max(1, text.count("\f") + 1),
        chunk_count=len(chunks),
        indexing_state="Embedded" if vector_result["embedded_count"] == len(chunks) and chunks else "No Text",
        extracted_text=text,
        metadata_json=json.dumps({
            "original_filename": file.filename,
            "sha256": hashlib.sha256(data).hexdigest(),
            "embedding": vector_result,
        }),
    )
    session.add(doc)
    session.commit()
    return document_payload(doc)


@router.get("/api/documents")
def list_documents(session: Session = Depends(get_session)) -> dict[str, Any]:
    return {"items": [document_payload(d) for d in session.scalars(select(Document).order_by(Document.upload_date.desc()))]}


@router.get("/api/documents/{document_id}/download")
def download_document(document_id: str, session: Session = Depends(get_session)) -> FileResponse:
    doc = session.get(Document, document_id)
    if not doc:
        raise HTTPException(404, "Unknown document.")
    path = app_path("APP_DOCUMENT_PATH", "./data/documents") / doc.file_name
    if not path.exists():
        raise HTTPException(404, "Stored document file is missing.")
    metadata = document_metadata(doc)
    filename = str(metadata.get("original_filename") or doc.file_name)
    return FileResponse(path, filename=filename, media_type="application/octet-stream")


@router.get("/api/documents/{document_id}/vectors")
def document_vectors(document_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    doc = session.get(Document, document_id)
    if not doc:
        raise HTTPException(404, "Unknown document.")
    rows = document_vector_rows(document_id)
    preview = [
        {
            "index": row.get("index"),
            "text": str(row.get("text", ""))[:900],
            "character_count": len(str(row.get("text", ""))),
            "has_embedding": isinstance(row.get("embedding"), list),
            "dimension": len(row.get("embedding")) if isinstance(row.get("embedding"), list) else 0,
        }
        for row in rows[:20]
    ]
    return {"document": document_payload(doc), "vector": vector_summary(document_id), "chunks": preview}


@router.post("/api/documents/{document_id}/reindex")
def reindex_document(document_id: str, session: Session = Depends(get_session)) -> dict[str, Any]:
    doc = session.get(Document, document_id)
    if not doc:
        raise HTTPException(404, "Unknown document.")
    chunks = chunk_text(doc.extracted_text or "")
    try:
        vector_result = write_vector_chunks(document_id, chunks, require_embeddings=bool(chunks))
    except RuntimeError as exc:
        doc.indexing_state = "Embedding Failed"
        doc.metadata_json = json.dumps({**document_metadata(doc), "embedding": {"error": str(exc), "chunk_count": len(chunks), "embedded_count": 0}})
        session.commit()
        raise HTTPException(502, f"Embedding indexing failed: {exc}") from exc
    doc.chunk_count = len(chunks)
    doc.indexing_state = "Embedded" if vector_result["embedded_count"] == len(chunks) and chunks else "No Text"
    doc.metadata_json = json.dumps({**document_metadata(doc), "embedding": vector_result})
    session.commit()
    return document_payload(doc)


@router.delete("/api/documents/{document_id}")
def delete_document(document_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    doc = session.get(Document, document_id)
    if not doc:
        raise HTTPException(404, "Unknown document.")
    app_path("APP_DOCUMENT_PATH", "./data/documents").joinpath(doc.file_name).unlink(missing_ok=True)
    document_vector_path(document_id).unlink(missing_ok=True)
    session.delete(doc)
    session.commit()
    return {"status": "deleted"}


@router.post("/api/chat")
def chat(payload: ChatIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    sources = retrieve_sources(payload, session)
    prompt_sources = [
        {
            "document_id": source["document_id"],
            "title": source["title"],
            "category": source["category"],
            "scope": source["scope"],
            "excerpt": source["excerpt"],
        }
        for source in sources
    ]
    db_context = ""
    if payload.conference_id:
        db_context = json.dumps(conference_payload(require_conference(session, payload.conference_id)), default=str)
    prompt = (
        f"Mode: {payload.mode}\nQuestion: {payload.question}\nConference facts: {db_context}\n"
        f"Retrieved document excerpts: {json.dumps(prompt_sources)}\n"
        "Answer directly using the retrieved excerpts as the RAG evidence. Cite document titles only, not chunk numbers. "
        "Distinguish rule from recommendation. State when information is unavailable in the uploaded knowledge base."
    )
    answer = call_azure_chat(prompt, system_prompt=assistant_system_prompt(session))
    if not answer:
        if sources:
            answer = "Based on the indexed local documents, the relevant excerpts are listed below. Azure OpenAI is not configured or did not respond, so this answer is a retrieval summary rather than a generated policy interpretation."
        else:
            answer = "No matching indexed document content was found. Upload IEEE or ITSS guidance documents, then ask again."
    return {"answer": answer, "sources": citation_sources(sources), "mode": payload.mode}


@router.post("/api/email-drafts", status_code=201)
def generate_email_draft(payload: EmailDraftIn, session: Session = Depends(get_session)) -> dict[str, Any]:
    conference = require_conference(session, payload.conference_id)
    issues = [require_issue(session, issue_id) for issue_id in payload.issue_ids]
    selected_contacts = [
        {
            "role": contact.role,
            "name": contact.name,
            "email": contact.email,
            "organization": contact.organization,
            "is_primary": contact.is_primary,
        }
        for contact in conference.contacts
        if contact.active
    ]
    facts = {
        "conference": {
            key: value
            for key, value in conference_payload(conference).items()
            if key
            in {
                "conference_number",
                "canonical_name",
                "official_title",
                "conference_series",
                "lifecycle_phase",
                "conference_status",
                "status_band",
                "score",
                "start_date",
                "end_date",
                "city",
                "country",
                "application_status",
                "mou_status",
                "finance_status",
                "publication_status",
                "last_source_update",
                "comments",
            }
        },
        "issues": [issue_payload(issue, conference) for issue in issues],
        "contacts": selected_contacts,
        "recipients": {
            "names": payload.recipient_names,
            "to": [str(address) for address in payload.recipient_addresses],
            "cc": [str(address) for address in payload.cc_addresses],
        },
        "purpose": payload.purpose,
        "optional_instructions": payload.instructions,
        "tone": payload.tone,
    }
    prompt = (
        "You are an IEEE ITSS conference operations email writer. Draft one ready-to-send email, not notes, not a summary, "
        "and not a restatement of the prompt. Use the selected conference, selected issues, recipients, purpose, tone, "
        "and optional instructions as drafting guidance. The optional_instructions field is instruction to you; do not copy it "
        "verbatim into the email and do not label it as context. Turn it into natural email wording only when relevant. "
        "Use only the supplied facts. Do not invent dates, approvals, names, or policy requirements. "
        "Write a clear subject and a polished body with greeting, concise context, requested action or decision, and professional closing. "
        "Return strict JSON only with keys: subject, body. "
        f"Facts: {json.dumps(facts, default=str)}"
    )
    generated = call_azure_chat(prompt)
    subject = f"{conference.canonical_name}: {payload.purpose}"
    body = local_email_body(conference, issues, payload)
    generator = "Local composer"
    if generated:
        parsed = parse_email_generation(generated)
        if parsed:
            subject = parsed.get("subject") or subject
            body = parsed.get("body") or body
            generator = "Azure OpenAI"
        else:
            body = generated
            generator = "Azure OpenAI"
    draft = EmailDraft(
        conference_id=conference.id,
        related_issues_json=json.dumps(payload.issue_ids),
        recipient_names=", ".join(payload.recipient_names),
        recipient_addresses=", ".join(str(v) for v in payload.recipient_addresses),
        cc_addresses=", ".join(str(v) for v in payload.cc_addresses),
        subject=subject,
        body=body,
        tone=payload.tone,
        generator=generator,
    )
    session.add(draft)
    session.commit()
    return email_payload(draft)


@router.get("/api/email-drafts")
def list_email_drafts(session: Session = Depends(get_session)) -> dict[str, Any]:
    return {"items": [email_payload(d) for d in session.scalars(select(EmailDraft).order_by(EmailDraft.created_at.desc()))]}


@router.delete("/api/email-drafts/{draft_id}")
def delete_email_draft(draft_id: str, session: Session = Depends(get_session)) -> dict[str, str]:
    draft = session.get(EmailDraft, draft_id)
    if not draft:
        raise HTTPException(404, "Unknown email draft.")
    session.delete(draft)
    session.commit()
    return {"status": "deleted"}


@router.get("/api/settings")
def settings(session: Session = Depends(get_session)) -> dict[str, Any]:
    score_weights = json.loads(session.get(DashboardSettings, "score_weights").value_json)
    scoring = score_settings(session)
    portfolio_start_year = int(json.loads(session.get(DashboardSettings, "portfolio_start_year").value_json))
    feature_setting = session.get(DashboardSettings, "feature_flags")
    feature_flags = {**default_feature_flags(), **(json.loads(feature_setting.value_json) if feature_setting else {})}
    role_setting = session.get(DashboardSettings, "role_permissions")
    role_permissions = merged_role_permissions(json.loads(role_setting.value_json) if role_setting else None)
    config = reference_config(session)
    mappings = {m.source_value: m.normalized_value for m in session.scalars(select(StatusMapping).where(StatusMapping.active.is_(True)))}
    return {
        "azure_openai": azure_status(mask=True),
        "embeddings": embedding_status(mask=True),
        "score_weights": score_weights,
        "score_settings": scoring,
        "portfolio_start_year": portfolio_start_year,
        "status_mappings": mappings,
        "feature_flags": feature_flags,
        "role_permissions": role_permissions,
        "assistant_system_prompt": assistant_system_prompt(session),
        "roles": role_catalog(),
        "permission_catalog": permission_catalog(),
        "reference_config": config,
        "reference_config_labels": REFERENCE_CONFIG_LABELS,
        "milestone_date_defaults": milestone_date_offsets(session),
    }


@router.patch("/api/settings")
def update_settings(payload: SettingsUpdate, session: Session = Depends(get_session)) -> dict[str, Any]:
    if payload.portfolio_start_year is not None:
        setting = session.get(DashboardSettings, "portfolio_start_year")
        setting.value_json = json.dumps(payload.portfolio_start_year)
    if payload.score_weights is not None:
        setting = session.get(DashboardSettings, "score_weights")
        setting.value_json = json.dumps(payload.score_weights)
    if payload.score_settings is not None:
        setting = session.get(DashboardSettings, "score_settings")
        formula = payload.score_settings.get("score_formula") if isinstance(payload.score_settings, dict) else None
        if formula is not None:
            try:
                validate_score_formula(str(formula))
            except ValueError as exc:
                raise HTTPException(422, str(exc)) from exc
        merged = merged_score_settings(payload.score_settings)
        if setting:
            setting.value_json = json.dumps(merged)
        else:
            session.add(DashboardSettings(key="score_settings", value_json=json.dumps(merged)))
        for conference in session.scalars(select(Conference)):
            recalculate(conference, session)
    if payload.status_mappings:
        for source, normalized in payload.status_mappings.items():
            mapping = session.scalar(select(StatusMapping).where(StatusMapping.source_value == source))
            if mapping:
                mapping.normalized_value = normalized
            else:
                session.add(StatusMapping(source_value=source, normalized_value=normalized))
    if payload.feature_flags is not None:
        setting = session.get(DashboardSettings, "feature_flags")
        merged = {**default_feature_flags(), **payload.feature_flags}
        if setting:
            setting.value_json = json.dumps(merged)
        else:
            session.add(DashboardSettings(key="feature_flags", value_json=json.dumps(merged)))
    if payload.role_permissions is not None:
        setting = session.get(DashboardSettings, "role_permissions")
        merged_roles = merged_role_permissions(payload.role_permissions)
        if setting:
            setting.value_json = json.dumps(merged_roles)
        else:
            session.add(DashboardSettings(key="role_permissions", value_json=json.dumps(merged_roles)))
    if payload.assistant_system_prompt is not None:
        prompt = " ".join(payload.assistant_system_prompt.split())
        setting = session.get(DashboardSettings, "assistant_system_prompt")
        if setting:
            setting.value_json = json.dumps(prompt)
        else:
            session.add(DashboardSettings(key="assistant_system_prompt", value_json=json.dumps(prompt)))
    session.commit()
    return settings(session)


@router.post("/api/settings/recalculate-scores")
def recalculate_all_scores(session: Session = Depends(get_session)) -> dict[str, Any]:
    updated = 0
    for conference in session.scalars(select(Conference)):
        recalculate(conference, session)
        updated += 1
    session.commit()
    return {"updated": updated, "message": f"Recalculated scores and statuses for {updated} conferences."}


@router.patch("/api/settings/reference-config")
def update_reference_config(payload: ReferenceConfigUpdate, session: Session = Depends(get_session)) -> dict[str, Any]:
    existing = reference_config(session)
    merged = {**existing, **payload.reference_config}
    clean = sanitize_reference_config(merged)
    setting = session.get(DashboardSettings, "reference_config")
    if setting:
        setting.value_json = json.dumps(clean)
    else:
        session.add(DashboardSettings(key="reference_config", value_json=json.dumps(clean)))
    changes = enforce_reference_config(session, clean)
    session.commit()
    return {**settings(session), "reference_cleanup": changes}


@router.post("/api/settings/verify-azure-openai")
def verify_azure() -> dict[str, Any]:
    status = azure_status(mask=True)
    checked_at = now().isoformat()
    if not status["configured"]:
        missing = [
            name
            for name, value in {
                "PLATFORM_LLM_BASE_URL or AZURE_OPENAI_ENDPOINT": llm_base_url(),
                "PLATFORM_LLM_API_KEY or AZURE_OPENAI_API_KEY": llm_api_key(),
                "AZURE_OPENAI_DEPLOYMENT or AZURE_OPENAI_CHAT_DEPLOYMENT": chat_deployment() if is_azure_llm_base_url(llm_base_url()) else "not required",
            }.items()
            if not value
        ]
        return {**status, "ok": False, "checked_at": checked_at, "message": "LLM gateway is not fully configured.", "missing": missing}
    try:
        content = llm_chat_completion_text(
            system_prompt="You are a concise connection test for the IEEE ITSS dashboard.",
            user_prompt="Reply with ok.",
            temperature=0,
            max_tokens=80,
            timeout_seconds=30,
        )
        return {**status, "ok": True, "checked_at": checked_at, "message": "LLM chat endpoint responded.", "sample": content}
    except Exception as exc:
        cause = repr(exc.__cause__) if getattr(exc, "__cause__", None) else ""
        return {
            **status,
            "ok": False,
            "checked_at": checked_at,
            "message": str(exc) or exc.__class__.__name__,
            "error_type": exc.__class__.__name__,
            "error_cause": cause,
            "deployment": chat_deployment() or default_model_alias(),
        }


@router.post("/api/settings/test-llm-message")
def test_llm_message(payload: LlmTestMessage) -> dict[str, Any]:
    status = azure_status(mask=True)
    checked_at = now().isoformat()
    if not status["configured"]:
        return {**status, "ok": False, "checked_at": checked_at, "message": "LLM gateway is not fully configured.", "response": ""}
    try:
        content = llm_chat_completion_text(
            system_prompt="You are a concise connection test for the IEEE ITSS dashboard.",
            user_prompt=payload.message,
            temperature=0,
            max_tokens=80,
            timeout_seconds=45,
        )
        return {
            **status,
            "ok": True,
            "checked_at": checked_at,
            "message": "LLM test message completed.",
            "response": content,
            "deployment": chat_deployment() or default_model_alias(),
        }
    except Exception as exc:
        cause = repr(exc.__cause__) if getattr(exc, "__cause__", None) else ""
        return {
            **status,
            "ok": False,
            "checked_at": checked_at,
            "message": str(exc) or exc.__class__.__name__,
            "response": "",
            "error_type": exc.__class__.__name__,
            "error_cause": cause,
            "deployment": chat_deployment() or default_model_alias(),
        }


@router.post("/api/settings/verify-embeddings")
def verify_embeddings() -> dict[str, Any]:
    load_local_env()
    status = embedding_status(mask=True)
    checked_at = now().isoformat()
    base = embedding_base_url().rstrip("/")
    if not base:
        return {**status, "ok": False, "checked_at": checked_at, "message": "TEI embedding endpoint is not configured."}
    try:
        health = httpx.get(f"{base}/health", timeout=10, verify=ssl_context())
        health.raise_for_status()
        info: dict[str, Any] = {}
        try:
            info_response = httpx.get(f"{base}/info", timeout=10, verify=ssl_context())
            if info_response.status_code == 200:
                info = info_response.json()
        except (httpx.RequestError, json.JSONDecodeError):
            info = {}
        sample = embed_texts(["IEEE ITSS embedding connectivity test"])
        dimension = len(sample[0]) if sample and sample[0] else 0
        return {
            **status,
            "ok": True,
            "checked_at": checked_at,
            "message": "TEI embedding endpoint is reachable.",
            "dimension": dimension,
            "model_info": info,
        }
    except Exception as exc:
        return {
            **status,
            "ok": False,
            "checked_at": checked_at,
            "message": "TEI embedding endpoint needs attention.",
            "error_type": type(exc).__name__,
            "error_cause": str(exc),
        }


@router.post("/api/settings/refresh-conference-facts")
def refresh_all_conference_facts(session: Session = Depends(get_session)) -> dict[str, Any]:
    conferences = list(session.scalars(select(Conference)))
    synced = 0
    for conference in conferences:
        ensure_milestones(conference, session)
        sync_conference_facts_from_milestones(conference)
        detect_issues(conference, session)
        recalculate(conference, session)
        create_snapshot(conference, session, reason="bulk refresh")
        synced += 1
    session.commit()
    return {"status": "ok", "synced": synced}


class MilestoneDatesPayload(BaseModel):
    milestone_date_defaults: dict[str, dict[str, Any]] | None = None


@router.post("/api/settings/recalculate-milestone-dates")
def recalculate_all_milestone_dates(payload: MilestoneDatesPayload | None = None, session: Session = Depends(get_session)) -> dict[str, Any]:
    if payload and payload.milestone_date_defaults:
        clean: dict[str, dict[str, Any]] = {}
        for code, offset in payload.milestone_date_defaults.items():
            anchor = str(offset.get("anchor", "start")).lower()
            months = int(offset.get("months", 0))
            days = int(offset.get("days", 0))
            clean[code.upper()] = {"anchor": anchor if anchor in {"start", "end"} else "start", "months": months, "days": days}
        # Merge with defaults to keep any unspecified codes
        merged = dict(MILESTONE_DATE_DEFAULTS)
        merged.update(clean)
        # Store in dashboard settings for persistence
        setting = session.get(DashboardSettings, "milestone_date_defaults")
        if setting:
            setting.value_json = json.dumps(merged)
        else:
            session.add(DashboardSettings(key="milestone_date_defaults", value_json=json.dumps(merged)))
        session.flush()
    conferences = list(session.scalars(select(Conference)))
    offsets = milestone_date_offsets(session)
    updated = 0
    for conference in conferences:
        for milestone in conference.milestones:
            code = milestone.definition.code
            new_due = milestone_due_date(code, conference.start_date, conference.end_date, offsets=offsets)
            if new_due and (milestone.due_date is None or milestone.due_date != new_due):
                milestone.due_date = new_due
                milestone.last_updated = now()
                updated += 1
        recalculate(conference, session)
    session.commit()
    # Return updated settings
    return {**settings(session), "updated": updated}


def require_conference(session: Session, conference_id: str) -> Conference:
    conference = session.get(Conference, conference_id)
    if not conference:
        raise HTTPException(404, "Unknown conference.")
    return conference


def require_issue(session: Session, issue_id: str) -> Issue:
    issue = session.get(Issue, issue_id)
    if not issue:
        raise HTTPException(404, "Unknown issue.")
    return issue


def create_snapshot(conference: Conference, session: Session, *, reason: str) -> Snapshot:
    payload = conference_payload(conference)
    payload["reason"] = reason
    snapshot = Snapshot(conference_id=conference.id, payload_json=json.dumps(payload, default=str))
    session.add(snapshot)
    return snapshot


CONFERENCE_COLUMNS = [
    "report_date",
    "conference_number",
    "acronym",
    "year",
    "official_title",
    "conference_series",
    "conference_category",
    "sponsorship_type",
    "lifecycle_phase",
    "conference_status",
    "start_date",
    "end_date",
    "submission_deadline",
    "notification_date",
    "camera_ready_deadline",
    "city",
    "state",
    "country",
    "ieee_region",
    "estimated_attendees",
    "actual_attendees",
    "website",
    "project_code",
    "application_status",
    "application_submitted_date",
    "application_approved_date",
    "mou_status",
    "mou_signed_date",
    "sponsor_summary",
    "sponsor_percentage",
    "application_comments",
    "finance_status",
    "finance_report_type",
    "finance_report_date",
    "currency",
    "total_income_current",
    "total_expense_current",
    "budgeted_income_total",
    "budgeted_expense_total",
    "bank_account_status",
    "payflow_status",
    "budget_approval_status",
    "accounting_close_date",
    "financial_required_items",
    "loan_status",
    "certificate_of_accuracy_status",
    "bank_closure_status",
    "tax_vat_status",
    "financial_comments",
    "publication_form_status",
    "publication_form_submitted_date",
    "loa_status",
    "loa_date",
    "publication_status",
    "media_type",
    "media_received_date",
    "xmldoc_date",
    "proceedings_submitted_date",
    "quality_review_status",
    "xplore_posting_date",
    "no_show_status",
    "publication_comments",
    "overall_comments",
]
CONTACT_COLUMNS = ["conference_number", "acronym", "year", "contact_role", "name", "email", "organization", "phone", "is_primary"]
ISSUE_COLUMNS = ["conference_number", "acronym", "year", "issue_key", "title", "description", "category", "severity", "issue_status", "review_assessment", "owner", "due_date", "comment"]
MILESTONE_COLUMNS = ["conference_number", "acronym", "year", "milestone_code", "milestone_name", "status", "due_date", "comments"]


def build_import_preview(filename: str, data: bytes, session: Session) -> dict[str, Any]:
    file_hash = hashlib.sha256(data).hexdigest()
    if session.scalar(select(ImportBatch).where(ImportBatch.file_hash == file_hash, ImportBatch.import_status == "Applied")):
        raise HTTPException(409, "This exact file was already imported.")
    rows = read_import_rows(filename, data)
    seen: set[tuple[str | None, str, int]] = set()
    results = []
    summary = {"rows": len(rows), "new": 0, "changed": 0, "unchanged": 0, "conflicts": 0}
    conflicts = []
    milestone_rows_raw = read_import_milestone_rows(filename, data)
    milestone_results = []
    for mrow in milestone_rows_raw:
        m_acronym = str(mrow.get("acronym") or "").strip().upper()
        m_year_raw = mrow.get("year")
        try:
            m_year = int(m_year_raw)
        except (TypeError, ValueError):
            m_year = 0
        m_conference_number = clean_cell(mrow.get("conference_number"))
        m_code = str(mrow.get("milestone_code") or "").strip().upper()
        m_errors = []
        if not m_code:
            m_errors.append("milestone_code is required")
        m_conference = match_conference(session, m_conference_number, m_acronym, m_year)
        if not m_conference:
            m_errors.append("Could not match conference")
        m_definition = None
        if m_code and m_conference:
            m_definition = session.scalar(select(MilestoneDefinition).where(MilestoneDefinition.code == m_code))
            if not m_definition:
                m_errors.append(f"Unknown milestone code: {m_code}")
        m_milestone = None
        if m_conference and m_definition:
            m_milestone = session.scalar(
                select(ConferenceMilestone).where(
                    ConferenceMilestone.conference_id == m_conference.id,
                    ConferenceMilestone.definition_id == m_definition.id,
                )
            )
        m_changes = []
        if m_milestone:
            for mfield in ["status", "due_date", "comments"]:
                if mfield not in mrow:
                    continue
                raw_val = mrow[mfield]
                if mfield == "status":
                    new_val = normalize_status(raw_val, session) if raw_val else None
                elif mfield == "due_date":
                    new_val = parse_date(raw_val) if raw_val else None
                else:
                    new_val = clean_cell(raw_val)
                if new_val is _SKIP_IMPORT_VALUE:
                    continue
                old_val = getattr(m_milestone, mfield, None)
                if comparable_import_value(old_val) != comparable_import_value(new_val):
                    m_changes.append({"field": mfield, "old": old_val, "new": new_val})
        milestone_results.append({
            "conference_number": m_conference_number or "",
            "acronym": m_acronym,
            "year": m_year,
            "milestone_code": m_code,
            "milestone_name": m_definition.name if m_definition else "",
            "matched": m_conference is not None and m_definition is not None,
            "errors": m_errors,
            "changes": m_changes,
            "source": mrow,
        })

    for index, row in enumerate(rows, start=2):
        errors: list[str] = []
        acronym = str(row.get("acronym") or "").strip()
        year_raw = row.get("year")
        try:
            year = int(year_raw)
        except (TypeError, ValueError):
            year = 0
            errors.append("year must be a number")
        conference_number = clean_cell(row.get("conference_number"))
        identity = (conference_number, normalize_acronym(acronym), year)
        if identity in seen:
            errors.append("duplicate row for the same conference")
        seen.add(identity)
        conference = match_conference(session, conference_number, acronym, year)
        if not conference and not minimum_import_fields(row):
            errors.append("new conferences require acronym, year, official_title, conference_series, sponsorship_type, and lifecycle_phase")
        changes = []
        if conference:
            for field, column in IMPORT_FIELD_MAP.items():
                if column not in row:
                    continue
                try:
                    new = import_value_for_field(field, row.get(column), session)
                except (TypeError, ValueError) as exc:
                    errors.append(f"{column}: {exc}")
                    continue
                if new is _SKIP_IMPORT_VALUE:
                    continue
                old = getattr(conference, field, None)
                if comparable_import_value(old) != comparable_import_value(new):
                    changes.append({"field": field, "old": old, "new": new})
            if changes:
                summary["changed"] += 1
            else:
                summary["unchanged"] += 1
        else:
            summary["new"] += 1
        validation = "valid" if not errors else "error"
        if errors:
            summary["conflicts"] += 1
            conflicts.append({"row_number": index, "errors": errors})
        results.append({"row_number": index, "matched_conference_id": conference.id if conference else None, "match_method": "number/acronym-year" if conference else "new", "validation_result": validation, "errors": errors, "changes": changes, "source": row})
    return {"file_name": filename, "file_hash": file_hash, "summary": summary, "rows": results, "conflicts": conflicts, "milestone_rows": milestone_results}


def read_import_rows(filename: str, data: bytes) -> list[dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix in {".xlsx", ".xlsm"}:
            workbook = pd.ExcelFile(io.BytesIO(data))
            if not workbook.sheet_names:
                raise HTTPException(400, "The workbook does not contain any worksheets.")
            sheet_name = "Conferences" if "Conferences" in workbook.sheet_names else workbook.sheet_names[0]
            frame = pd.read_excel(workbook, sheet_name=sheet_name)
        elif suffix == ".csv":
            frame = pd.read_csv(io.BytesIO(data))
        else:
            raise HTTPException(400, "Only .xlsx, .xlsm, and .csv files are supported.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(400, f"Could not read import file: {exc}") from exc
    frame.columns = [str(column).strip() for column in frame.columns]
    if frame.empty:
        raise HTTPException(400, "The import file does not contain any data rows.")
    missing = [column for column in ["acronym", "year", "official_title"] if column not in frame.columns]
    if missing:
        raise HTTPException(400, "Missing required columns: " + ", ".join(missing))
    return frame.fillna("").to_dict(orient="records")


def read_import_milestone_rows(filename: str, data: bytes) -> list[dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    if suffix not in {".xlsx", ".xlsm"}:
        return []
    try:
        workbook = pd.ExcelFile(io.BytesIO(data))
        if "Milestones" not in workbook.sheet_names:
            return []
        frame = pd.read_excel(workbook, sheet_name="Milestones")
    except Exception:
        return []
    if frame.empty:
        return []
    frame.columns = [str(column).strip() for column in frame.columns]
    missing = [c for c in ["acronym", "year", "milestone_code"] if c not in frame.columns]
    if missing:
        return []
    return frame.fillna("").to_dict(orient="records")


IMPORT_FIELD_MAP = {
    "conference_number": "conference_number",
    "official_title": "official_title",
    "conference_series": "conference_series",
    "conference_category": "conference_category",
    "sponsorship_type": "sponsorship_type",
    "lifecycle_phase": "lifecycle_phase",
    "conference_status": "conference_status",
    "start_date": "start_date",
    "end_date": "end_date",
    "submission_deadline": "submission_deadline",
    "notification_date": "notification_date",
    "camera_ready_deadline": "camera_ready_deadline",
    "city": "city",
    "state_province": "state",
    "country": "country",
    "ieee_region": "ieee_region",
    "website": "website",
    "estimated_attendees": "estimated_attendees",
    "actual_attendees": "actual_attendees",
    "project_code": "project_code",
    "application_status": "application_status",
    "application_submitted_date": "application_submitted_date",
    "application_approved_date": "application_approved_date",
    "mou_status": "mou_status",
    "mou_signed_date": "mou_signed_date",
    "finance_status": "finance_status",
    "currency": "currency",
    "total_income_current": "total_income_current",
    "total_expense_current": "total_expense_current",
    "budgeted_income_total": "budgeted_income_total",
    "budgeted_expense_total": "budgeted_expense_total",
    "accounting_close_date": "accounting_close_date",
    "publication_status": "publication_status",
    "proceedings_submitted_date": "proceedings_submitted_date",
    "xplore_posting_date": "xplore_posting_date",
    "comments": "overall_comments",
}

STATUS_IMPORT_FIELDS = {"application_status", "mou_status", "finance_status", "publication_status", "conference_status"}
DATE_IMPORT_FIELDS = {
    "start_date",
    "end_date",
    "submission_deadline",
    "notification_date",
    "camera_ready_deadline",
    "application_submitted_date",
    "application_approved_date",
    "mou_signed_date",
    "accounting_close_date",
    "proceedings_submitted_date",
    "xplore_posting_date",
}
INTEGER_IMPORT_FIELDS = {"estimated_attendees", "actual_attendees"}
MONEY_IMPORT_FIELDS = {"total_income_current", "total_expense_current", "budgeted_income_total", "budgeted_expense_total"}
MILESTONE_FIELDS = {"status", "due_date", "comments"}
_SKIP_IMPORT_VALUE = object()


def match_conference(session: Session, conference_number: str | None, acronym: str, year: int) -> Conference | None:
    if conference_number:
        found = session.scalar(select(Conference).where(Conference.conference_number == conference_number))
        if found:
            return found
    if acronym and year:
        return session.scalar(select(Conference).where(Conference.normalized_acronym == normalize_acronym(acronym), Conference.year == year, Conference.parent_conference_id.is_(None)))
    return None


def minimum_import_fields(row: dict[str, Any]) -> bool:
    return all(clean_cell(row.get(field)) for field in ["acronym", "year", "official_title", "conference_series", "sponsorship_type", "lifecycle_phase"])


def conference_export_row(conference: Conference) -> dict[str, Any]:
    row = {column: "" for column in CONFERENCE_COLUMNS}
    for field, column in IMPORT_FIELD_MAP.items():
        value = getattr(conference, field, None)
        row[column] = value.isoformat() if isinstance(value, date) else value
    row.update(
        {
            "report_date": date.today().isoformat(),
            "acronym": conference.acronym,
            "year": conference.year,
            "official_title": conference.official_title,
            "conference_number": normalize_record_number(conference.conference_number),
            "application_comments": conference.application_status_raw or "",
            "financial_comments": "",
            "publication_comments": "",
        }
    )
    return row


def selected_import_fields(selected: Any, row: dict[str, Any]) -> set[str] | None:
    if selected is None:
        return set(IMPORT_FIELD_MAP.keys())
    row_number = row["row_number"]
    if isinstance(selected, list):
        return set(IMPORT_FIELD_MAP.keys()) if row_number in selected or str(row_number) in selected else None
    if isinstance(selected, dict):
        fields = selected.get(str(row_number), selected.get(row_number))
        if fields is None:
            return None
        if fields == "__all__":
            return set(IMPORT_FIELD_MAP.keys())
        if isinstance(fields, list):
            return {str(field) for field in fields if str(field) in IMPORT_FIELD_MAP}
    return None


def can_apply_partial_import_row(row: dict[str, Any], selected_fields: set[str]) -> bool:
    if not row.get("matched_conference_id") or not selected_fields:
        return False
    errored_fields = errored_import_fields(row.get("errors", []))
    if errored_fields is None:
        return False
    return selected_fields.isdisjoint(errored_fields)


def errored_import_fields(errors: list[str]) -> set[str] | None:
    by_column = {column: field for field, column in IMPORT_FIELD_MAP.items()}
    fields: set[str] = set()
    for error in errors:
        column, separator, _message = str(error).partition(":")
        if not separator:
            return None
        field = by_column.get(column.strip())
        if not field:
            return None
        fields.add(field)
    return fields


def apply_import_row(row: dict[str, Any], session: Session, batch: ImportBatch, *, selected_fields: set[str] | None = None) -> Conference:
    acronym = str(row["acronym"]).strip().upper()
    year = int(row["year"])
    conference = match_conference(session, clean_cell(row.get("conference_number")), acronym, year)
    fields_to_apply = selected_fields if selected_fields is not None else set(IMPORT_FIELD_MAP.keys())
    if conference is None:
        conference = Conference(
            conference_number=clean_cell(row.get("conference_number")),
            acronym=acronym,
            normalized_acronym=normalize_acronym(acronym),
            year=year,
            official_title=str(row["official_title"]).strip(),
            canonical_name=f"{acronym} {year}",
            conference_series=configured_value(session, "conference_series", str(row["conference_series"]).strip()),
            conference_category=clean_cell(row.get("conference_category")) or "Portfolio",
            sponsorship_type=str(row["sponsorship_type"]).strip(),
            lifecycle_phase=str(row["lifecycle_phase"]).strip(),
            phase_override=True,
        )
        session.add(conference)
        session.flush()
    for field, column in IMPORT_FIELD_MAP.items():
        if field not in fields_to_apply:
            continue
        value = import_value_for_field(field, row.get(column), session)
        if value is _SKIP_IMPORT_VALUE:
            continue
        old = getattr(conference, field)
        if comparable_import_value(old) != comparable_import_value(value):
            setattr(conference, field, value)
            session.add(FieldChange(conference_id=conference.id, field_name=field, old_value=str(old) if old else None, new_value=str(value) if value else None, change_type="Import", source=batch.original_filename, import_batch_id=batch.id))
    return conference


def import_value_for_field(field: str, raw: Any, session: Session) -> Any:
    value = clean_cell(raw)
    if value is None:
        return _SKIP_IMPORT_VALUE
    if value == "[CLEAR]":
        return None
    if field in STATUS_IMPORT_FIELDS:
        return normalize_status(value, session)
    if field == "lifecycle_phase":
        return configured_value(session, "lifecycle_phases", value)
    if field == "sponsorship_type":
        return configured_value(session, "sponsorship_types", value)
    if field == "conference_series":
        return configured_value(session, "conference_series", value)
    if field in DATE_IMPORT_FIELDS:
        return parse_date(value)
    if field in INTEGER_IMPORT_FIELDS:
        return int(float(value))
    if field in MONEY_IMPORT_FIELDS:
        return float(str(value).replace(",", ""))
    return value


def comparable_import_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def clean_cell(value: Any) -> str | None:
    if value is None or pd.isna(value):
        return None
    text = str(value).strip()
    return text or None


def safe_filename(filename: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]", "_", Path(filename).name)


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as exc:
            return f"Extraction failed: {exc}"
    if suffix == ".pdf":
        try:
            import fitz

            with fitz.open(stream=data, filetype="pdf") as doc:
                return "\f".join(page.get_text() for page in doc)
        except Exception as exc:
            return f"Extraction failed: {exc}"
    return data.decode("utf-8", errors="replace")


def chunk_text(text: str) -> list[str]:
    size = int(os.environ.get("RAG_CHUNK_SIZE", "1400"))
    overlap = int(os.environ.get("RAG_CHUNK_OVERLAP", "220"))
    normalized = clean_extracted_text(text)
    if not normalized:
        return []
    units = [unit.strip() for unit in re.split(r"\n{2,}|(?<=[.!?])\s+(?=[A-Z0-9])", normalized) if unit.strip()]
    chunks: list[str] = []
    current = ""
    for unit in units:
        if len(unit) > size:
            if current:
                chunks.append(current.strip())
                current = ""
            cursor = 0
            while cursor < len(unit):
                chunk = unit[cursor : cursor + size].strip()
                if chunk:
                    chunks.append(chunk)
                cursor += max(1, size - overlap)
            continue
        candidate = f"{current}\n\n{unit}".strip() if current else unit
        if len(candidate) <= size:
            current = candidate
        else:
            chunks.append(current.strip())
            previous_tail = current[-overlap:].strip() if overlap > 0 else ""
            current = f"{previous_tail}\n\n{unit}".strip() if previous_tail else unit
    if current:
        chunks.append(current.strip())
    return chunks


def clean_extracted_text(text: str) -> str:
    cleaned = text.replace("\r\n", "\n")
    cleaned = re.sub(r"[ƒ\u0192][\"“”]?", '"', cleaned)
    cleaned = re.sub(r'""([^"\n]{1,160})""', r'"\1"', cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


def write_vector_chunks(document_id: str, chunks: list[str], *, require_embeddings: bool = False) -> dict[str, Any]:
    path = app_path("APP_VECTOR_PATH", "./data/vector_store") / f"{document_id}.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    embeddings: list[list[float] | None] = [None for _ in chunks]
    error = ""
    if chunks and not embedding_configured():
        error = "TEI embedding endpoint is not configured."
    elif chunks:
        try:
            embeddings = embed_texts(chunks)
        except RuntimeError as exc:
            error = str(exc)
    if require_embeddings and chunks and (error or any(not embedding for embedding in embeddings)):
        raise RuntimeError(error or "TEI embedding endpoint did not return embeddings for every chunk.")
    dimension = len(next((embedding for embedding in embeddings if embedding), []))
    rows = [
        {
            "index": i,
            "text": chunk,
            "embedding": embeddings[i] if i < len(embeddings) else None,
            "metadata": {
                "document_id": document_id,
                "chunk_index": i,
                "character_count": len(chunk),
                "embedding_model": embedding_model(),
                "embedding_provider": "IAV on-prem TEI",
                "created_at": now().isoformat(),
            },
        }
        for i, chunk in enumerate(chunks)
    ]
    path.write_text(json.dumps(rows, indent=2), encoding="utf-8")
    embedded_count = sum(1 for row in rows if row.get("embedding"))
    return {
        "embedded_count": embedded_count,
        "chunk_count": len(chunks),
        "dimension": dimension,
        "model": embedding_model(),
        "provider": "IAV on-prem TEI",
        "vector_path": str(path),
        "error": error,
    }


def embedding_base_url() -> str:
    load_local_env()
    return (
        os.environ.get("TEI_EMBEDDING_BASE_URL")
        or os.environ.get("IAV_TEI_EMBEDDING_BASE_URL")
        or os.environ.get("EMBEDDING_BASE_URL")
        or ""
    ).strip()


def embedding_model() -> str:
    return (os.environ.get("TEI_EMBEDDING_MODEL") or os.environ.get("EMBEDDING_MODEL") or "qwen3").strip()


def embedding_api_key() -> str:
    return (os.environ.get("TEI_EMBEDDING_API_KEY") or os.environ.get("EMBEDDING_API_KEY") or "").strip()


def embedding_url() -> str:
    base = embedding_base_url().rstrip("/")
    if not base:
        return ""
    return base if base.endswith("/v1/embeddings") else f"{base}/v1/embeddings"


def embedding_configured() -> bool:
    return bool(embedding_base_url())


def embedding_status(*, mask: bool) -> dict[str, Any]:
    base = embedding_base_url()
    return {
        "configured": bool(base),
        "provider": "IAV on-prem TEI",
        "endpoint": mask_endpoint(base) if mask else base,
        "route": "/v1/embeddings",
        "model": embedding_model(),
        "api_key_required": False,
        "api_key_present": bool(embedding_api_key()),
    }


def embed_texts(texts: list[str]) -> list[list[float] | None]:
    clean_texts = [text.strip() for text in texts]
    if not clean_texts:
        return []
    url = embedding_url()
    if not url:
        raise RuntimeError("TEI_EMBEDDING_BASE_URL is required for embeddings.")
    batch_size = int(os.environ.get("TEI_EMBEDDING_BATCH_SIZE", "16"))
    timeout = float(os.environ.get("TEI_EMBEDDING_TIMEOUT_SECONDS", "30"))
    vectors: list[list[float] | None] = []
    for start in range(0, len(clean_texts), max(1, batch_size)):
        batch = clean_texts[start : start + max(1, batch_size)]
        payload: dict[str, Any] = {"input": batch, "encoding_format": "float"}
        model = embedding_model()
        if model:
            payload["model"] = model
        headers = {"Content-Type": "application/json"}
        api_key = embedding_api_key()
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        try:
            response = httpx.post(url, json=payload, headers=headers, timeout=timeout, verify=ssl_context())
            response.raise_for_status()
            body = response.json()
        except httpx.HTTPStatusError as exc:
            raise RuntimeError(provider_error_message(exc.response)) from exc
        except (httpx.RequestError, json.JSONDecodeError) as exc:
            raise RuntimeError(f"Embedding request failed: {exc}") from exc
        data = body.get("data") if isinstance(body, dict) else None
        if not isinstance(data, list):
            raise RuntimeError("Embedding response did not include a data list.")
        by_index: dict[int, list[float]] = {}
        for item in data:
            if isinstance(item, dict) and isinstance(item.get("index"), int) and isinstance(item.get("embedding"), list):
                by_index[item["index"]] = [float(value) for value in item["embedding"]]
        vectors.extend([by_index.get(index) for index in range(len(batch))])
    return vectors


def template_type_from_filename(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".doc", ".docx"}:
        return "Word"
    if suffix in {".xls", ".xlsx", ".xlsm", ".csv"}:
        return "Excel"
    if suffix == ".pdf":
        return "PDF"
    if suffix in {".ppt", ".pptx"}:
        return "PowerPoint"
    if suffix in {".txt", ".md"}:
        return "Text"
    return suffix.lstrip(".").upper() or "File"


def template_payload(template: TemplateFile) -> dict[str, Any]:
    return {
        "id": template.id,
        "template_name": template.template_name,
        "short_description": template.short_description,
        "category": template.category,
        "template_type": template.template_type,
        "original_filename": template.original_filename,
        "last_update": template.updated_at.isoformat(),
        "upload_date": template.upload_date.isoformat(),
    }


def cosine_similarity(a: list[float], b: list[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(y * y for y in b))
    if not norm_a or not norm_b:
        return 0.0
    return dot / (norm_a * norm_b)


def document_payload(doc: Document) -> dict[str, Any]:
    metadata = document_metadata(doc)
    embedding = metadata.get("embedding") if isinstance(metadata, dict) else {}
    return {
        "id": doc.id,
        "title": doc.title,
        "file_name": doc.file_name,
        "document_category": doc.document_category,
        "knowledge_scope": doc.knowledge_scope,
        "conference_id": doc.conference_id,
        "conference_series": doc.conference_series,
        "version": doc.version,
        "source_url": doc.source_url,
        "upload_date": doc.upload_date.isoformat(),
        "active": doc.active,
        "extraction_state": doc.extraction_state,
        "indexing_state": doc.indexing_state,
        "page_count": doc.page_count,
        "chunk_count": doc.chunk_count,
        "embedding": embedding if isinstance(embedding, dict) else {},
    }


def document_vector_path(document_id: str) -> Path:
    return app_path("APP_VECTOR_PATH", "./data/vector_store") / f"{document_id}.json"


def document_vector_rows(document_id: str) -> list[dict[str, Any]]:
    path = document_vector_path(document_id)
    if not path.exists():
        return []
    try:
        rows = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    return [row for row in rows if isinstance(row, dict)] if isinstance(rows, list) else []


def vector_summary(document_id: str) -> dict[str, Any]:
    path = document_vector_path(document_id)
    rows = document_vector_rows(document_id)
    embedded = [row for row in rows if isinstance(row.get("embedding"), list)]
    first_embedding = embedded[0].get("embedding") if embedded else []
    return {
        "exists": path.exists(),
        "path": str(path),
        "chunk_count": len(rows),
        "embedded_count": len(embedded),
        "dimension": len(first_embedding) if isinstance(first_embedding, list) else 0,
        "model": embedding_model(),
        "provider": "IAV on-prem TEI",
        "updated_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat() if path.exists() else None,
    }


def document_metadata(doc: Document) -> dict[str, Any]:
    try:
        metadata = json.loads(doc.metadata_json or "{}")
    except json.JSONDecodeError:
        metadata = {}
    return metadata if isinstance(metadata, dict) else {}


def retrieve_sources(payload: ChatIn, session: Session) -> list[dict[str, Any]]:
    terms = [term.lower() for term in re.findall(r"[A-Za-z0-9]{4,}", payload.question)]
    documents = list(session.scalars(select(Document).where(Document.active.is_(True))))
    scored: list[tuple[float, Document, str, str, int | None]] = []
    query_embedding: list[float] | None = None
    requested_scope = (payload.knowledge_scope or "IEEE ITSS").strip()
    search_all_scopes = requested_scope in {"All KBs", "All"}
    if embedding_configured():
        try:
            embeddings = embed_texts([payload.question])
            query_embedding = embeddings[0] if embeddings else None
        except RuntimeError:
            query_embedding = None
    for doc in documents:
        if not search_all_scopes and requested_scope and requested_scope != doc.knowledge_scope:
            continue
        vector_path = app_path("APP_VECTOR_PATH", "./data/vector_store") / f"{doc.id}.json"
        if query_embedding and vector_path.exists():
            try:
                vector_rows = json.loads(vector_path.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                vector_rows = []
            vector_hits = 0
            for row in vector_rows if isinstance(vector_rows, list) else []:
                embedding = row.get("embedding") if isinstance(row, dict) else None
                text = row.get("text", "") if isinstance(row, dict) else ""
                if isinstance(embedding, list) and text:
                    score = cosine_similarity(query_embedding, [float(value) for value in embedding])
                    if payload.conference_id and doc.conference_id == payload.conference_id:
                        score += 0.05
                    if payload.conference_series and doc.conference_series == payload.conference_series:
                        score += 0.03
                    scored.append((score, doc, text[:900], "vector", row.get("index") if isinstance(row.get("index"), int) else None))
                    vector_hits += 1
            if vector_hits:
                continue
        text = doc.extracted_text or ""
        lower = text.lower()
        score = float(sum(lower.count(term) for term in terms))
        if payload.conference_id and doc.conference_id == payload.conference_id:
            score += 5
        if payload.conference_series and doc.conference_series == payload.conference_series:
            score += 3
        if score:
            excerpt_start = min([lower.find(term) for term in terms if lower.find(term) >= 0] or [0])
            excerpt = text[excerpt_start : excerpt_start + 700]
            scored.append((score, doc, excerpt, "keyword", None))
    scored.sort(key=lambda item: item[0], reverse=True)
    top_k = int(os.environ.get("RAG_TOP_K", "8"))
    return [
        {
            "document_id": doc.id,
            "title": doc.title,
            "category": doc.document_category,
            "scope": doc.knowledge_scope,
            "excerpt": excerpt,
            "score": score,
            "retrieval": retrieval,
            "chunk_index": chunk_index,
        }
        for score, doc, excerpt, retrieval, chunk_index in scored[:top_k]
    ]


def citation_sources(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    citations: dict[str, dict[str, Any]] = {}
    for source in sources:
        document_id = str(source.get("document_id") or "")
        if not document_id or document_id in citations:
            continue
        citations[document_id] = {
            "document_id": document_id,
            "title": source.get("title"),
            "category": source.get("category"),
            "scope": source.get("scope"),
        }
    return list(citations.values())


def azure_status(*, mask: bool) -> dict[str, Any]:
    endpoint = llm_base_url()
    api_key = llm_api_key()
    version = azure_api_version()
    chat = chat_deployment()
    embedding = os.environ.get("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "")
    ca_bundle = azure_ca_bundle()
    proxy = azure_proxy()
    model_alias = default_model_alias()
    is_azure = is_azure_llm_base_url(endpoint)
    return {
        "configured": bool(endpoint and api_key and (chat or not is_azure)),
        "provider": "Azure OpenAI" if is_azure else "OpenAI-compatible gateway",
        "endpoint": endpoint if not mask else mask_endpoint(endpoint),
        "api_version": version,
        "chat_deployment": chat or model_alias,
        "model_alias": model_alias,
        "embedding_deployment": embedding,
        "api_key_present": bool(api_key),
        "ca_bundle_configured": bool(ca_bundle),
        "ca_bundle": Path(ca_bundle).name if mask and ca_bundle else ca_bundle,
        "proxy_configured": bool(proxy),
        "proxy": mask_proxy(proxy) if mask else proxy,
    }


def mask_endpoint(endpoint: str) -> str:
    if not endpoint:
        return ""
    return re.sub(r"^(https://[^./]+).*", r"\1...", endpoint)


def chat_deployment() -> str:
    return os.environ.get("AZURE_OPENAI_CHAT_DEPLOYMENT") or os.environ.get("AZURE_OPENAI_DEPLOYMENT", "")


def llm_base_url() -> str:
    return (os.environ.get("PLATFORM_LLM_BASE_URL") or os.environ.get("AZURE_OPENAI_ENDPOINT", "")).strip()


def llm_api_key() -> str:
    return (os.environ.get("PLATFORM_LLM_API_KEY") or os.environ.get("AZURE_OPENAI_API_KEY", "")).strip()


def allowed_model_aliases() -> tuple[str, ...]:
    aliases = tuple(
        item.strip()
        for item in os.environ.get("PLATFORM_ALLOWED_MODEL_ALIASES", "default-reasoning").split(",")
        if item.strip()
    )
    return aliases or ("default-reasoning",)


def default_model_alias() -> str:
    return allowed_model_aliases()[0]


def azure_api_version() -> str:
    return os.environ.get("AZURE_OPENAI_API_VERSION", "2024-12-01-preview").strip()


def is_azure_llm_base_url(base_url: str) -> bool:
    return "openai.azure.com" in base_url.lower()


def llm_chat_url(base_url: str, model_alias: str | None = None) -> str:
    base = base_url.rstrip("/")
    if is_azure_llm_base_url(base):
        deployment = chat_deployment() or model_alias or default_model_alias()
        if "/openai/deployments/" in base:
            url = base if base.endswith("/chat/completions") else f"{base}/chat/completions"
        else:
            url = f"{base}/openai/deployments/{deployment}/chat/completions"
        separator = "&" if "?" in url else "?"
        return url if "api-version=" in url else f"{url}{separator}api-version={azure_api_version()}"
    return base if base.endswith("/chat/completions") else f"{base}/chat/completions"


def azure_client() -> AzureOpenAI:
    timeout = float(os.environ.get("AZURE_OPENAI_TIMEOUT_SECONDS", "20"))
    proxy = azure_proxy()
    return AzureOpenAI(
        azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
        api_key=os.environ["AZURE_OPENAI_API_KEY"],
        api_version=os.environ["AZURE_OPENAI_API_VERSION"],
        http_client=httpx.Client(verify=ssl_context(), timeout=timeout, proxy=proxy or None),
    )


def ssl_context() -> ssl.SSLContext:
    ca_bundle = azure_ca_bundle()
    context = ssl.create_default_context(cafile=ca_bundle if ca_bundle else None)
    if os.name != "nt" or os.environ.get("AZURE_OPENAI_USE_WINDOWS_CERT_STORE", "1") in {"0", "false", "False"}:
        return context
    if not hasattr(ssl, "enum_certificates"):
        return context
    loaded = 0
    for store_name in ("ROOT", "CA"):
        try:
            certificates = ssl.enum_certificates(store_name)
        except OSError:
            continue
        for certificate, encoding, _trust in certificates:
            if encoding != "x509_asn":
                continue
            try:
                context.load_verify_locations(cadata=ssl.DER_cert_to_PEM_cert(certificate))
                loaded += 1
            except ssl.SSLError:
                continue
    context._windows_certificates_loaded = loaded  # type: ignore[attr-defined]
    return context


def azure_ca_bundle() -> str:
    for key in ("AZURE_OPENAI_CA_BUNDLE", "SSL_CERT_FILE", "REQUESTS_CA_BUNDLE"):
        value = os.environ.get(key, "").strip()
        if value:
            return value
    return ""


def azure_proxy() -> str:
    return os.environ.get("AZURE_OPENAI_PROXY", "").strip()


def mask_proxy(proxy: str) -> str:
    if not proxy:
        return ""
    return re.sub(r"//([^:@/]+):([^@/]+)@", r"//\1:***@", proxy)


def llm_chat_completion_text(
    *,
    system_prompt: str,
    user_prompt: str,
    model_alias: str | None = None,
    temperature: float | None = None,
    max_tokens: int | None = None,
    timeout_seconds: float | None = None,
    response_format: dict[str, str] | None = None,
) -> str:
    body = llm_chat_completion_body(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        model_alias=model_alias,
        temperature=temperature,
        max_tokens=max_tokens,
        timeout_seconds=timeout_seconds,
        response_format=response_format,
    )
    try:
        return str(body["choices"][0]["message"]["content"])
    except (KeyError, IndexError, TypeError) as exc:
        raise RuntimeError("LLM response did not contain message content.") from exc


def llm_chat_completion_body(
    *,
    system_prompt: str,
    user_prompt: str,
    model_alias: str | None = None,
    temperature: float | None = None,
    max_tokens: int | None = None,
    timeout_seconds: float | None = None,
    response_format: dict[str, str] | None = None,
) -> dict[str, Any]:
    base_url = llm_base_url()
    api_key = llm_api_key()
    if not base_url or not api_key:
        raise RuntimeError("PLATFORM_LLM_BASE_URL and PLATFORM_LLM_API_KEY are required.")
    model = model_alias or default_model_alias()
    is_azure = is_azure_llm_base_url(base_url)
    if is_azure and not chat_deployment() and not model:
        raise RuntimeError("Azure OpenAI endpoint is configured, but AZURE_OPENAI_DEPLOYMENT is missing.")
    url = llm_chat_url(base_url, model)
    payload: dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    if temperature is not None and supports_temperature(model):
        payload["temperature"] = temperature
    if max_tokens is not None:
        token_key = "max_completion_tokens" if uses_max_completion_tokens(model) else "max_tokens"
        payload[token_key] = max_tokens
    if response_format is not None:
        payload["response_format"] = response_format
    timeout = timeout_seconds if timeout_seconds is not None else float(os.environ.get("AZURE_OPENAI_TIMEOUT_SECONDS", "45"))
    try:
        response = httpx.post(
            url,
            json=payload,
            headers={
                **({"api-key": api_key} if is_azure else {"Authorization": f"Bearer {api_key}"}),
                "Content-Type": "application/json",
            },
            timeout=timeout,
            verify=ssl_context(),
            proxy=azure_proxy() or None,
        )
        response.raise_for_status()
        body = response.json()
        if not isinstance(body, dict):
            raise RuntimeError("LLM response JSON must be an object.")
        return body
    except httpx.HTTPStatusError as exc:
        raise RuntimeError(provider_error_message(exc.response)) from exc
    except httpx.RequestError as exc:
        if os.name == "nt":
            try:
                return post_chat_completion_with_powershell(
                    url=url,
                    payload=payload,
                    api_key=api_key,
                    is_azure=is_azure,
                    timeout_seconds=timeout,
                )
            except RuntimeError as fallback_exc:
                raise RuntimeError(str(fallback_exc)) from exc
        raise RuntimeError(f"LLM request failed: {exc}") from exc


def effective_model_name(model_alias: str) -> str:
    return chat_deployment() or model_alias


def uses_max_completion_tokens(model_alias: str) -> bool:
    model = effective_model_name(model_alias).lower()
    return model.startswith("gpt-5") or model.startswith("o1") or model.startswith("o3") or model.startswith("o4")


def supports_temperature(model_alias: str) -> bool:
    model = effective_model_name(model_alias).lower()
    return not (model.startswith("gpt-5") or model.startswith("o1") or model.startswith("o3") or model.startswith("o4"))


def post_chat_completion_with_powershell(
    *,
    url: str,
    payload: dict[str, Any],
    api_key: str,
    is_azure: bool,
    timeout_seconds: float,
) -> dict[str, Any]:
    script = r"""
$body = [Console]::In.ReadToEnd()
$headers = @{ "Content-Type" = "application/json" }
if ($env:LLM_IS_AZURE -eq "true") {
  $headers["api-key"] = $env:LLM_API_KEY
} else {
  $headers["Authorization"] = "Bearer " + $env:LLM_API_KEY
}
try {
  $response = Invoke-WebRequest `
    -Uri $env:LLM_URL `
    -Method Post `
    -UseBasicParsing `
    -ContentType "application/json" `
    -Headers $headers `
    -Body $body
  [Console]::Out.WriteLine([int]$response.StatusCode)
  [Console]::Out.Write($response.Content)
} catch {
  if ($_.Exception.Response -ne $null) {
    $statusCode = [int]$_.Exception.Response.StatusCode
    $stream = $_.Exception.Response.GetResponseStream()
    $reader = New-Object System.IO.StreamReader($stream)
    $content = $reader.ReadToEnd()
    if ([string]::IsNullOrWhiteSpace($content) -and $_.ErrorDetails -ne $null) {
      $content = $_.ErrorDetails.Message
    }
    [Console]::Out.WriteLine($statusCode)
    [Console]::Out.Write($content)
  } else {
    [Console]::Error.WriteLine($_.Exception.Message)
    exit 1
  }
}
"""
    env = {
        **os.environ,
        "LLM_API_KEY": api_key,
        "LLM_IS_AZURE": "true" if is_azure else "false",
        "LLM_URL": url,
    }
    try:
        completed = subprocess.run(
            ["powershell.exe", "-NoProfile", "-NonInteractive", "-Command", script],
            input=json.dumps(payload),
            text=True,
            capture_output=True,
            env=env,
            timeout=timeout_seconds + 5,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise RuntimeError(f"PowerShell LLM fallback failed: {exc}") from exc
    if completed.returncode != 0:
        raise RuntimeError(f"PowerShell LLM fallback failed: {completed.stderr.strip() or completed.stdout.strip()}")
    status_line, _, response_body = completed.stdout.lstrip().partition("\n")
    try:
        status_code = int(status_line.strip())
    except ValueError as exc:
        raise RuntimeError("PowerShell LLM fallback returned invalid output.") from exc
    try:
        body = json.loads(response_body)
    except json.JSONDecodeError as exc:
        raise RuntimeError("LLM response was not valid JSON.") from exc
    if status_code >= 400:
        raise RuntimeError(provider_error_payload_message(status_code, body))
    if not isinstance(body, dict):
        raise RuntimeError("LLM response JSON must be an object.")
    return body


def provider_error_message(response: httpx.Response) -> str:
    try:
        payload = response.json()
    except json.JSONDecodeError:
        return f"LLM request failed with status {response.status_code}."
    return provider_error_payload_message(response.status_code, payload)


def provider_error_payload_message(status_code: int, payload: Any) -> str:
    if isinstance(payload, dict):
        error = payload.get("error")
        if isinstance(error, dict):
            message = error.get("message")
            if isinstance(message, str) and message.strip():
                return f"LLM rejected the request: {message}"
        detail = payload.get("detail")
        if isinstance(detail, str) and detail.strip():
            return f"LLM rejected the request: {detail}"
    return f"LLM request failed with status {status_code}."


def call_azure_chat(prompt: str, *, system_prompt: str | None = None) -> str | None:
    status = azure_status(mask=False)
    if not status["configured"]:
        return None
    try:
        return llm_chat_completion_text(
            system_prompt=system_prompt or "You assist IEEE ITSS conference operations. Do not invent policy or missing facts.",
            user_prompt=prompt,
            temperature=float(os.environ.get("AZURE_OPENAI_TEMPERATURE", "0.2")),
            max_tokens=int(os.environ.get("AZURE_OPENAI_MAX_TOKENS", "2000")),
        )
    except Exception:
        return None


def local_recommendation(conference: Conference, issue: Issue) -> str:
    return (
        f"For {conference.canonical_name}, review the source field for '{issue.title}', assign an owner, "
        "confirm the latest status with the responsible IEEE or conference contact, and update the issue assessment. "
        "No AI policy interpretation was generated because Azure OpenAI is unavailable."
    )


def local_email_body(conference: Conference, issues: list[Issue], payload: EmailDraftIn) -> str:
    recipient = payload.recipient_names[0] if payload.recipient_names else ""
    greeting = f"Dear {recipient}," if recipient else "Dear colleagues,"
    location = ", ".join(part for part in (conference.city, conference.country) if part)
    date_window = " to ".join(part.isoformat() for part in (conference.start_date, conference.end_date) if part)
    conference_bits = [conference.canonical_name]
    if conference.conference_number:
        conference_bits.append(f"IEEE Conference Record Number {conference.conference_number}")
    if location:
        conference_bits.append(location)
    if date_window:
        conference_bits.append(date_window)
    context = "; ".join(conference_bits)
    issue_lines = [
        f"- {issue.title}: {issue.description or issue.category}. Current assessment: {issue.review_assessment}; severity: {issue.severity}."
        for issue in issues
    ]
    issue_section = "\n".join(issue_lines) if issue_lines else "- Please confirm whether there are any open items that need IEEE ITSS attention."
    instruction_action = instruction_to_request(payload.instructions)
    action_lines = [
        f"Could you please provide a short update for {context}?",
        "Please include the current status, responsible owner, expected completion date, and any support needed from IEEE ITSS.",
    ]
    if instruction_action:
        action_lines.append(instruction_action)
    return (
        f"{greeting}\n\n"
        f"I am following up regarding {context}.\n\n"
        f"The purpose of this note is to {payload.purpose.lower()}.\n\n"
        f"The items I would like to confirm are:\n{issue_section}\n\n"
        f"{' '.join(action_lines)}\n\n"
        "Thank you,\n"
        "IEEE ITSS Conferences"
    )


def instruction_to_request(instructions: str | None) -> str:
    if not instructions:
        return ""
    text = " ".join(instructions.strip().split())
    if not text:
        return ""
    text = text[0].lower() + text[1:] if len(text) > 1 else text.lower()
    if text.endswith((".", "?", "!")):
        text = text[:-1]
    return f"Please also {text}."


def parse_email_generation(generated: str) -> dict[str, str] | None:
    text = generated.strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return None
    if not isinstance(parsed, dict):
        return None
    subject = parsed.get("subject")
    body = parsed.get("body")
    if not isinstance(subject, str) and not isinstance(body, str):
        return None
    return {"subject": subject if isinstance(subject, str) else "", "body": body if isinstance(body, str) else ""}


def parse_generated_issues(generated: str, session: Session) -> list[GeneratedIssue]:
    text = generated.strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as exc:
        raise HTTPException(502, "LLM issue response was not valid JSON.") from exc
    if not isinstance(parsed, dict):
        raise HTTPException(502, "LLM issue response must be a JSON object.")
    raw_issues = parsed.get("issues", [])
    if not isinstance(raw_issues, list):
        raise HTTPException(502, "LLM issue response must include an issues list.")
    allowed_categories = set(allowed_reference_values(session, "issue_categories"))
    allowed_severities = set(allowed_reference_values(session, "issue_severities"))
    allowed_assessments = set(allowed_reference_values(session, "review_assessments"))
    issues: list[GeneratedIssue] = []
    for raw in raw_issues[:20]:
        if not isinstance(raw, dict):
            continue
        title = " ".join(str(raw.get("title", "")).strip().split())
        if not title:
            continue
        category = str(raw.get("category") or "AI Review")
        severity = str(raw.get("severity") or "Medium")
        assessment = str(raw.get("review_assessment") or "Needs Follow-up")
        try:
            due_date = parse_date(raw.get("due_date"))
        except (TypeError, ValueError):
            due_date = None
        issues.append(
            GeneratedIssue(
                title=title[:240],
                description=str(raw.get("description") or ""),
                category=category if category in allowed_categories else "AI Review",
                severity=severity if severity in allowed_severities else "Medium",
                review_assessment=assessment if assessment in allowed_assessments else "Needs Follow-up",
                owner=str(raw.get("owner") or "").strip() or None,
                due_date=due_date,
            )
        )
    return issues


def email_payload(draft: EmailDraft) -> dict[str, Any]:
    return {
        "id": draft.id,
        "conference_id": draft.conference_id,
        "related_issues": json.loads(draft.related_issues_json),
        "recipient_names": draft.recipient_names,
        "recipient_addresses": draft.recipient_addresses,
        "cc_addresses": draft.cc_addresses,
        "subject": draft.subject,
        "body": draft.body,
        "tone": draft.tone,
        "generator": draft.generator,
        "created_at": draft.created_at.isoformat(),
        "saved": draft.saved,
    }


def simple_pdf(lines: list[str]) -> bytes:
    escaped_lines = [line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)") for line in lines]
    text_commands = ["BT", "/F1 12 Tf", "50 750 Td", "14 TL"]
    for index, line in enumerate(escaped_lines[:48]):
        if index:
            text_commands.append("T*")
        text_commands.append(f"({line}) Tj")
    text_commands.append("ET")
    stream = "\n".join(text_commands).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    output = io.BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(output.tell())
        output.write(f"{number} 0 obj\n".encode("ascii"))
        output.write(body)
        output.write(b"\nendobj\n")
    xref = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    return output.getvalue()


def current_surplus(conference: Conference) -> float | None:
    if conference.total_income_current is None or conference.total_expense_current is None:
        return None
    return conference.total_income_current - conference.total_expense_current


def sqlite_health() -> dict[str, Any]:
    db_path = app_path("APP_DATABASE_PATH", "./data/itss_dashboard.db")
    with sqlite3.connect(db_path) as connection:
        row = connection.execute("select count(*) from conferences").fetchone()
    return {"database": str(db_path), "conference_count": row[0]}
